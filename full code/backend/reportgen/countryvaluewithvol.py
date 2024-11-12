import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx.shared import Inches,Cm
import copy
import os
from docx.shared import RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE
import datetime
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, Border, Side, NamedStyle,Alignment
from openpyxl.cell.cell import WriteOnlyCell
from openpyxl.worksheet.table import Table, TableStyleInfo
from mpmath import *                            


def auto100(regen):
    regend=[]
    regen=[mpf(str(i).replace("%","")) for i in regen]
    sm=sum(regen)
    
    if sm!=100.0:
        if sm>100.0:
            diff=mpf((sm-100.0)/len(regen))
            for i in regen:
                regend.append(i-diff)
        else :
            diff=mpf((100.0-sm)/len(regen))
            for i in regen:
                regend.append(i+diff)
    else:
        regend=regen
    return(regend)


def toc100(toc):
    # for i in toc:
    for k in toc: 
        sttc2=[]
        endtc2=[]
        sttc2vol=[]
        endtc2vol=[]
        for tc2 in k["toc_component"]:
            sttc2.append(tc2["split"][0])
            endtc2.append(tc2["split"][1])
            sttc2vol.append(tc2["vsplit"][0])
            endtc2vol.append(tc2["vsplit"][1])
            if tc2["toc2_component"]!=[]:
                sttc3=[]
                endtc3=[]
                sttc3vol=[]
                endtc3vol=[]
                for tc3 in tc2["toc2_component"]:
                    sttc3.append(tc3["split"][0])
                    endtc3.append(tc3["split"][1])
                    sttc3vol.append(tc3["vsplit"][0])
                    endtc3vol.append(tc3["vsplit"][1])
                    if tc3["toc3_component"]!=[]:
                        sttc4=[]
                        endtc4=[]
                        sttc4vol=[]
                        endtc4vol=[]
                        for tc4 in tc3["toc3_component"]:
                            # print(tc4)
                            sttc4.append(tc4["split"][0])
                            endtc4.append(tc4["split"][1])
                            sttc4vol.append(tc4["vsplit"][0])
                            endtc4vol.append(tc4["vsplit"][1])
                        
                        
                        regst13=auto100(sttc4)
                        regen13=auto100(endtc4)
                        regst13vol=auto100(sttc4vol)
                        regen13vol=auto100(endtc4vol)
                        for tc4 in range(len(tc3["toc3_component"])):
                            tc3["toc3_component"][tc4]["split"][0]=regst13[tc4]
                            tc3["toc3_component"][tc4]["split"][1]=regen13[tc4]
                            tc3["toc3_component"][tc4]["vsplit"][0]=regst13vol[tc4]
                            tc3["toc3_component"][tc4]["vsplit"][1]=regen13vol[tc4]
                        

                regst12=auto100(sttc3)
                regen12=auto100(endtc3)
                regst12vol=auto100(sttc3vol)
                regen12vol=auto100(endtc3vol)
                for tc3 in range(len(tc2["toc2_component"])):
                    tc2["toc2_component"][tc3]["split"][0]=regst12[tc3]
                    tc2["toc2_component"][tc3]["split"][1]=regen12[tc3]
                    tc2["toc2_component"][tc3]["vsplit"][0]=regst12vol[tc3]
                    tc2["toc2_component"][tc3]["vsplit"][1]=regen12vol[tc3]


        regst1=auto100(sttc2)
        regen1=auto100(endtc2)  
        regst1vol=auto100(sttc2vol)
        regen1vol=auto100(endtc2vol)  
        for tc2 in range(len(k["toc_component"])):
            k["toc_component"][tc2]["split"][0]=regst1[tc2]
            k["toc_component"][tc2]["split"][1]=regen1[tc2]
            k["toc_component"][tc2]["vsplit"][0]=regst1vol[tc2]
            k["toc_component"][tc2]["vsplit"][1]=regen1vol[tc2]
            







def tocl2vol(toc,year,yearvol,g):
    for j in range(len(toc)):
        for i in toc[j]["toc_component"]:
            s1=mpf((str(i["split"][0])).replace("%",""))
            s2=mpf((str(i["split"][1])).replace("%",""))
            s1v=mpf((str(i["vsplit"][0])).replace("%",""))
            s2v=mpf((str(i["vsplit"][1])).replace("%",""))
            step=(s2-s1)/(g[1]-g[0])
            stepv=(s2v-s1v)/(g[1]-g[0])
            i["split"]={}
            i["vsplit"]={}
            for k in range (g[0],g[1]+1):
                v=(s1*mpf(str(year[str(k)]).replace(",","")))/100
                vl=(s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100
                i["split"][str(k)]=float(v)
                i["vsplit"][str(k)]=float(vl)
                s1+=step
                s1v+=stepv
    return toc


def tocl3vol(toc,g):
    # print(toc)
    for j in range(len(toc)):
        for c in toc[j]["toc_component"]:
            year=c["split"]
            yearvol=c["vsplit"]
            # print(c["toc2_name"])
            for t in (c["toc2_component"]):
                s1=mpf((str(t["split"][0])).replace("%",""))
                s2=mpf((str(t["split"][1])).replace("%",""))
                s1v=mpf((str(t["vsplit"][0])).replace("%",""))
                s2v=mpf((str(t["vsplit"][1])).replace("%",""))
                stepv=(s2v-s1v)/(g[1]-g[0])
                step=(s2-s1)/(g[1]-g[0])
                t["split"]={}
                t["vsplit"]={}
                for k in range (g[0],(g[1]+1)):
                    t["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
                    t["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
                    s1+=step
                    s1v+=stepv
    return toc 

            
                
def tocl4vol(toc,g):
    for j in range(len(toc)):
        for c in toc[j]["toc_component"]:
            for t in (c["toc2_component"]):
                year=t["split"]
                yearvol=t["vsplit"]
                # print(t["toc3_name"])
                for r in (t["toc3_component"]):
                    s1=mpf((str(r["split"][0])).replace("%",""))
                    s2=mpf((str(r["split"][1])).replace("%",""))
                    s1v=mpf((str(r["vsplit"][0])).replace("%",""))
                    s2v=mpf((str(r["vsplit"][1])).replace("%",""))
                    stepv=(s2v-s1v)/(g[1]-g[0])
                    step=(s2-s1)/(g[1]-g[0])
                    r["split"]={}
                    r["vsplit"]={}
                    for k in range (g[0],(g[1]+1)):
                        r["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
                        r["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
                        s1+=step
                        s1v+=stepv
    return toc


def abc(toc,reportname,country):
    pdlist={}
    count=0
    count1=0
    count2=0
    for i in toc:
        pdlist[country+" "+reportname+" MARKET BY, "+i["toc_name"]+"#"+str(count)]={"heading":1,"sin":0,"comp":i["toc_name"],"data":{},"vdata":{},"val":country+" "+reportname+" MARKET VALUE BY, "+i["toc_name"],"vol":country+" "+reportname+" MARKET VOLUME BY, "+i["toc_name"],"name":reportname}
        for l2 in i["toc_component"]:
            pdlist[country+" "+reportname+" MARKET BY, "+i["toc_name"]+"#"+str(count)]["data"][l2["toc2_name"]]=l2["split"]
            pdlist[country+" "+reportname+" MARKET BY, "+i["toc_name"]+"#"+str(count)]["vdata"][l2["toc2_name"]]=l2["vsplit"]
            count1=count
            if l2["toc2_component"]!=[]:
                pdlist[country+" "+l2["toc2_name"]+" MARKET BY, "+i["toc_name"]+"#"+str(count1)]={"heading":2,"sin":0,"comp":l2["toc2_name"],"data":{},"vdata":{},"val":country+" "+l2["toc2_name"]+" MARKET VALUE BY, "+i["toc_name"],"vol":country+" "+l2["toc2_name"]+" MARKET VOLUME BY, "+i["toc_name"] ,"name":l2["toc2_name"]}
                for l3 in l2["toc2_component"]:
                    pdlist[country+" "+l2["toc2_name"]+" MARKET BY, "+i["toc_name"]+"#"+str(count1)]["data"][l3["toc3_name"]]=l3["split"]
                    pdlist[country+" "+l2["toc2_name"]+" MARKET BY, "+i["toc_name"]+"#"+str(count1)]["vdata"][l3["toc3_name"]]=l3["vsplit"]
                    count2=count1
                    if l3["toc3_component"]!=[]:
                        pdlist[country+" "+l3["toc3_name"]+" MARKET BY, "+i["toc_name"]+"#"+str(count2)]={"heading":3,"sin":0,"comp":l3["toc3_name"],"data":{},"vdata":{},"val":country+" "+l3["toc3_name"]+" MARKET VALUE BY, "+i["toc_name"],"vol":country+" "+l3["toc3_name"]+" MARKET VOLUME BY, "+i["toc_name"],"name":l3["toc3_name"]}
                        for l4 in l3["toc3_component"]:
                            pdlist[country+" "+l3["toc3_name"]+" MARKET BY, "+i["toc_name"]+"#"+str(count2)]["data"][l4["toc4_name"]]=l4["split"]
                            pdlist[country+" "+l3["toc3_name"]+" MARKET BY, "+i["toc_name"]+"#"+str(count2)]["vdata"][l4["toc4_name"]]=l4["vsplit"]
                            pdlist[country+" "+l4["toc4_name"]+" MARKET BY, "+i["toc_name"]+"#"+str(count2)]={"heading":4,"sin":1,"comp":"Market","data":{l4["toc4_name"]:l4["split"]},"vdata":{l4["toc4_name"]:l4["vsplit"]},"val":country+" "+l4["toc4_name"]+" MARKET VALUE BY, "+i["toc_name"],"vol":country+" "+l4["toc4_name"]+" MARKET VOLUME BY, "+i["toc_name"],"name":l4["toc4_name"]}
                    
                    else:
                        pdlist[country+" "+l3["toc3_name"]+" MARKET BY,"+i["toc_name"]+"#"+str(count2)]={"heading":3,"sin":1,"comp":"Market","data":{l3["toc3_name"]:l3["split"]},"vdata":{l3["toc3_name"]:l3["vsplit"]},"val":country+" "+l3["toc3_name"]+" MARKET VALUE BY, "+i["toc_name"],"vol":country+" "+l3["toc3_name"]+" MARKET VOLUME BY, "+i["toc_name"],"name":l3["toc3_name"]}    
                    count2+=1    
            else:
                pdlist[country+" "+l2["toc2_name"]+" MARKET BY,"+i["toc_name"]+"#"+str(count)]={"heading":2,"sin":1,"comp":"Market","data":{l2["toc2_name"]:l2["split"]},"vdata":{l2["toc2_name"]:l2["vsplit"]},"val":country+" "+l2["toc2_name"]+" MARKET VALUE BY, "+i["toc_name"],"vol":country+" "+l2["toc2_name"]+" MARKET VOLUME BY, "+i["toc_name"],"name":l2["toc2_name"]}

            count1+=1
        count+=1
    return pdlist
                   

def plot(j,t):
    dir=os.getcwd()
    if j.shape[1]==1:
        ax=j.plot(legend=None,kind="bar",figsize=(8,6.25))
    else:
        ax=j.plot(kind="bar",figsize=(8,6.25))
        plt.legend(loc="upper left")
    ax.set_axisbelow(True)
    ax.yaxis.grid(color='gray', linestyle='dashed')
    plt.xticks(rotation='horizontal')
    # plt.legend(loc="lower center", bbox_to_anchor=(0.5, -0.2),fancybox=True, shadow=True,ncol=2,handletextpad=0.5)
    plt.savefig(dir+"/temp/"+t+'example.png',bbox_inches='tight',pad_inches =0.05)
    plt.close()
    return(dir+"/temp/"+t+'example.png')



def parastyle(headno):
    if headno==1:
        return "Main_Heading-NM"
    elif headno==2:
        return "Head2_NM"
    elif headno==3 :
        return "Head3_NM"
    elif headno==4:
        return "head4_NM"


def writeup(df,name,market,value):    
    value=value.replace(" USD","")
    if df.shape[0]==2:
        ml=("The ",market," ",df.index[0]," Market was valued at $",str(round(df[df.columns[0]][0],2))," ",value," in ",df.columns[0],", and is projected to reach $",str(round(df[df.columns[-2]][0],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]][0],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". ")
        return("".join(ml))
    
    elif df.shape[0]==3:
        ml=["The ",market," ",name," Market was valued at $",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach $",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with $",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach $",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        b=df1[df1.columns[-1]].idxmax()
        ml+=[b," segment is anticipated to reach $",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
        return("".join(ml))
    else:
        ml=["The ",market," ",name," Market was valued at $",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach $",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with $",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach $",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        b=df1[df1.columns[-1]].idxmax()
        ml+=[b," segment is anticipated to reach $",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%."]
        df2=df1.drop(index=[a,b])
        c=df2[df2.columns[0]].idxmax()
        d=df2[df2.columns[-1]].idxmax()
        e=round(((df[df.columns[0]][a]+df[df.columns[0]][c])/df[df.columns[0]]["Total"])*100,1)
        ml+=[a," and ",c," segments collectively expected to account for about ",str(e),"% share of the ",market," ",name," market in ",df.columns[0],", with the former constituting around ",str(round((df[df.columns[0]][a]/df[df.columns[0]]["Total"])*100,1)),"% share. ",b," and ",d," segments are expected to witness significant growth rates at a CAGR of ",str(round(df[df.columns[-1]][b],2)),"% and ",str(round(df[df.columns[-1]][d],2)),"% respectively, during the forecast period. Presently, share of these two segments is estimated to be around ",str(round(((df[df.columns[0]][b]+df[df.columns[0]][d])/df[df.columns[0]]["Total"])*100,2)),"% in the overall ",market," ",name," market in ",df.columns[0],", and is anticipated to reach ",str(round(((df[df.columns[-2]][b]+df[df.columns[-2]][d])/df[df.columns[-2]]["Total"])*100,2)),"% by ",df.columns[-2],". "]
        return("".join(ml))
   
    
def writeupvol(df,name,market,value):    
    if df.shape[0]==2:
        ml=("The total demand for ",market," ",df.index[0]," market was ",str(round(df[df.columns[0]][0],2))," ",value," in ",df.columns[0],", and is projected to reach ",str(round(df[df.columns[-2]][0],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]][0],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". ")
        return("".join(ml))
    
    elif df.shape[0]==3:
        ml=["The total demand for ",market," ",name," market was ",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach ",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with ",str(round(df[df.columns[0]][a],2))," in ",df.columns[0],", and is anticipated to reach ",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        return("".join(ml))
    else:
        ml=["The total demand for ",market," ",name," market was ",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach ",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with ",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach ",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        b=df1[df1.columns[-1]].idxmax()
        ml+=[b," segment is anticipated to reach ",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
        df2=df1.drop(index=[a,b])
        c=df2[df2.columns[0]].idxmax()
        d=df2[df2.columns[-1]].idxmax()
        e=round(((df[df.columns[0]][a]+df[df.columns[0]][c])/df[df.columns[0]]["Total"])*100,1)
        ml+=[a," and ",c," segments collectively expected to account for about ",str(e),"% share of the ",market," ",name," market in ",df.columns[0],", with the former constituting around ",str(round((df[df.columns[0]][a]/df[df.columns[0]]["Total"])*100,1)),"% share. ",b," and ",d," segments are expected to witness significant growth rates at a CAGR of ",str(round(df[df.columns[-1]][b],2)),"% and ",str(round(df[df.columns[-1]][d],2)),"% respectively, during the forecast period. Presently, share of these two segments is estimated to be around ",str(round(((df[df.columns[0]][b]+df[df.columns[0]][d])/df[df.columns[0]]["Total"])*100,2)),"% in the overall ",market," ",name," market in ",df.columns[0],", and is anticipated to reach ",str(round(((df[df.columns[-2]][b]+df[df.columns[-2]][d])/df[df.columns[-2]]["Total"])*100,2)),"% by ",df.columns[-2],". "]
        return("".join(ml))


def tocspl(toc,start,end):
    for l1 in toc:
        for l2 in l1["toc_component"]:
            for k1 in range (start,end):
                    l2["split"].pop(str(k1))
                    l2["vsplit"].pop(str(k1))
            if l2["toc2_component"]!=[]:
                for l3 in l2["toc2_component"]:
                    for k1 in range  (start,end):
                        l3["split"].pop(str(k1))
                        l3["vsplit"].pop(str(k1))
                    if l3["toc3_component"]!=[]:
                        for l4 in l3["toc3_component"]:
                            for k1 in range  (start,end):
                                l4["split"].pop(str(k1))
                                l4["vsplit"].pop(str(k1))



def createreport(pdlist,reportname,market,value,volume,tablesplit,g):
    dir =os.getcwd()
    doc = docx.Document("Template NM _New.docx")
    # reportname=reportname.split("(1)")[0]
    paragraph_format1 = doc.styles["Table1_Header"].paragraph_format
    paragraph_format2 = doc.styles['Table1_Left'].paragraph_format
    paragraph_format3 = doc.styles['Table1_Right'].paragraph_format
    paragraph_format4 = doc.styles['Table_No_NM'].paragraph_format
    paragraph_format5 = doc.styles['FIG_no_NM'].paragraph_format
    paragraph_format6 = doc.styles['Main_Heading-NM'].paragraph_format
    paragraph_format7 = doc.styles['Head2_NM'].paragraph_format
    paragraph_format8 = doc.styles['Head3_NM'].paragraph_format
    paragraph_format9 = doc.styles['head4_NM'].paragraph_format
    # paragraph_format10 = doc.styles['Normal'].paragraph_format
    # paragraph_format11 = doc.styles['Source_NM'].paragraph_format
    paragraph_format1.keep_with_next = True
    paragraph_format2.keep_with_next = True
    paragraph_format3.keep_with_next = True
    paragraph_format4.keep_with_next = True
    paragraph_format5.keep_with_next = True
    paragraph_format6.keep_with_next = True
    paragraph_format7.keep_with_next = True
    paragraph_format8.keep_with_next = True
    paragraph_format9.keep_with_next = True
    # paragraph_format10.keep_with_next = True
    # paragraph_format11.space_before=Pt(0)
    dt = datetime.datetime.now()
    t = str(datetime.datetime.timestamp(dt))

    Tabno = NamedStyle(name="Tab_no")
    Tabno.font = Font(name=' Franklin Gothic Medium Cond',
                    size=12,
                    bold=True,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')

    heading1 = NamedStyle(name="heading1")
    heading1.font = Font(name='Franklin Gothic Demi Cond',
                    size=19,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')


    heading2= NamedStyle(name="heading2")
    heading2.font = Font(name='Franklin Gothic Medium Cond',
                    size=14,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')


    source=NamedStyle(name="source")
    source.font=Font(name='Franklin Gothic Book',
                    size=7.5,
                    bold=False,
                    italic=True,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='545454')

    main_heading = NamedStyle(name="main_heading")
    main_heading.font = Font(name='Franklin Gothic Medium Cond',
                    size=35,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')

    number=NamedStyle(name="number",number_format="0.00")



    wb = Workbook(write_only=True)
    ws = wb.create_sheet()
    ws.sheet_view.showGridLines = False

    cellmain = WriteOnlyCell(ws)
    cellmain.style = main_heading
    cellmain.alignment = Alignment(horizontal='center', vertical='center')  

    cell = WriteOnlyCell(ws)
    cell.style = Tabno

    cellh1 = WriteOnlyCell(ws)
    cellh1.style = heading1
        

    cellh2 = WriteOnlyCell(ws)
    cellh2.style = heading2

        
    cell1 = WriteOnlyCell(ws)
    cell1.style = source

    cellnum= WriteOnlyCell(ws)
    cellnum.style = number

    def format_row(row, cell):
        for c in row:
            cell.value = c
            yield cell

    styletable = TableStyleInfo(name="TableStyleMedium2", showFirstColumn=True,
                       showLastColumn=True, showRowStripes=True, showColumnStripes=False)


    rnum =1      
    count=1
    count1=1
    count2=1
    tabname=1
    lstcolmain=chr(ord('c')+(g[1]-g[0]+2))
    mainr=[None,None,market+" "+reportname]
    mainrow=format_row(mainr, cellmain)
    ws.append(mainrow)
    ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcolmain+str(rnum))
    rnum+=1
    col=[str(i) for i in range(g[0],g[1]+1)]





    for k in pdlist:
        try:
            f=pd.DataFrame.from_dict(pdlist[k]["data"],orient='index')
            f = f.reindex(columns=col)
            fvol=pd.DataFrame.from_dict(pdlist[k]["vdata"],orient='index')
            fvol = fvol.reindex(columns=col)
            mt=(pdlist[k]["val"],", ",f.columns[0],"-",f.columns[-1],", ",value)
            mtv=(pdlist[k]["vol"],", ",f.columns[0],"-",f.columns[-1],", ",volume)
            doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))

            if (pdlist[k]["sin"])==0:
                if pdlist[k]["heading"]==1:
                    doc.add_paragraph("Overview",style=parastyle(2))
                
                f.loc["Total"] = f.sum()
                f['CAGR % '+"("+f.columns[1]+"-"+f.columns[-1]+")" ] = ((f.iloc[:, -1].div(f.iloc[:, 1]).pow(1./(len(f.columns) - 2)).sub(1))*100 )  
                fvol.loc["Total"] = fvol.sum()
                fvol['CAGR % '+"("+fvol.columns[1]+"-"+fvol.columns[-1]+")" ] = ((fvol.iloc[:, -1].div(fvol.iloc[:, 1]).pow(1./(len(fvol.columns) - 2)).sub(1))*100 ) 
                if volume=="Units":
                    fvol = fvol.round().astype({i:'int' for i in col })

                lstcol=chr(ord('c')+f.shape[1])
                if pdlist[k]["heading"]==1:
                    h1=[None,None,str(count1)+". "+(k.split("#")[0]).upper()]
                    count1+=1
                    count2=1
                    head_row = format_row(h1, cellh1)
                    ws.append(head_row)
                    ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                    rnum +=1  

                elif pdlist[k]["heading"]==2 :
                    h2=[None,None,str(count2)+". "+(k.split("#")[0]).upper()]
                    count2+=1
                    head2_row = format_row(h2, cellh2)
                    ws.append(head2_row)
                    ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                    rnum +=1  

                            
                r1=[None,None,"TABLE "+str(count)+". "+("".join(mt)).upper()]
                count+=1
                row1=[None,None] 
                first_row = format_row(r1, cell)
                ws.append(first_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1  
                rowtable=rnum
                df=list(dataframe_to_rows(f, index=True, header=True))
                for i in range (len(df)):
                    row1=[None,None]
                    row= row1+df[i]
                #         print((row))
                    if i==0:
                        row=list(map(str, df[i]))
                        row[0]=pdlist[k]["comp"]
                        row=row1+row
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  
                    elif i==1:
                        continue
                    else:
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  

                tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
                tabname+=1
                headings = [pdlist[k]["comp"]]+list(f.columns)
                tab1._initialise_columns()
                for column, value1 in zip(tab1.tableColumns, headings):
                    column.name = value1
            #     tab1.headerRowCount = False
                
                tab1.tableStyleInfo = styletable
                ws.add_table(tab1)
                sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
                sourcerow = format_row(sourcer, cell1)
                ws.append(sourcerow)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                ws.append([None])
                ws.append([None])
                rnum +=3


                r1=[None,None,"TABLE "+str(count)+". "+("".join(mtv)).upper()]
                count+=1
                row1=[None,None] 
                first_row = format_row(r1, cell)
                ws.append(first_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1 
                rowtable=rnum 
                df=list(dataframe_to_rows(fvol, index=True, header=True))
                for i in range (len(df)):
                    row1=[None,None]
                    row= row1+df[i]
                #         print((row))
                    if i==0:
                        row=list(map(str, df[i]))
                        row[0]=pdlist[k]["comp"]
                        row=row1+row
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  
                    elif i==1:
                        continue
                    else:
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  

                tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
                tabname+=1
                headings = [pdlist[k]["comp"]]+list(f.columns)
                tab1._initialise_columns()
                for column, value1 in zip(tab1.tableColumns, headings):
                    column.name = value1
            #     tab1.headerRowCount = False
                
                tab1.tableStyleInfo = styletable
                ws.add_table(tab1)
                sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
                sourcerow = format_row(sourcer, cell1)
                ws.append(sourcerow)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                ws.append([None])
                ws.append([None])
                rnum +=3





                if (tablesplit):
                    f=f.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)
                    fvol=fvol.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)

                f1=f.drop([f.columns[-1]],axis=1)
                f1=f1.drop(index=("Total"))
                p=doc.add_paragraph("".join(mt),style="FIG_no_NM")
                pic=plot(f1.T,t)
                ps=doc.add_picture(pic,width=Cm(17.2), height=Cm(8.51))
                os.remove(pic)
                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")
                # doc.add_paragraph(" ")


                
            
                p=doc.add_paragraph("".join(mt),style="Table_No_NM")
                table = doc.add_table(rows=f.shape[0]+1, cols=f.shape[1]+1)
                table.style = 'Table_Style_NM'
                table_cells = table._cells
                for r in range (len(table.rows)):
                    row=table.rows[r]
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                    if r==0:
                        row.height = Cm(1.1)
                    else:
                        row.height = Cm(0.7)
                
                table.cell(0,0).paragraphs[0].style="Table1_Header"
                table.cell(0,0).width = 1242000.2592
                table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
                table.columns[0].width= Cm(3.45)
                table.columns[-1].width=Cm(1.97)


                for k1 in range(1,len(f.columns)+1):
                    table.cell(0,k1).paragraphs[0].style="Table1_Header"
                    table.cell(0,k1).paragraphs[0].text=(f.columns[k1-1])
                for i in range(1,f.shape[0]+1):
                    table_cells[ i*(f.shape[1]+1)].paragraphs[0].style="Table1_Left"
                    table_cells[ i*(f.shape[1]+1)].paragraphs[0].text=(f.index[i-1])
                    for j in range(1,f.shape[1]+1):
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].style="Table1_Right"
                        if (j %(f.shape[1]))==0:
                            table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                            table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2))+"%")
                            
                        elif (j %(f.shape[1]+1))==0:
                            table_cells[j + i * (f.shape[1]+1)].width = 1242000.2592
                            table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))
                            
                        else:
                            table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                            table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))
                            

                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis ",style="Source_NM")

                doc.add_paragraph(writeup(f,pdlist[k]["name"],market,value),style="Normal")


                f1vol=fvol.drop([f.columns[-1]],axis=1)
                f1vol=f1vol.drop(index=("Total"))
                p=doc.add_paragraph("".join(mtv),style="FIG_no_NM")
                picv=plot(f1vol.T,t)
                ps=doc.add_picture(picv,width=Cm(17.2), height=Cm(8.51))
                os.remove(picv)
                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")

               


                p=doc.add_paragraph("".join(mtv),style="Table_No_NM")
                table = doc.add_table(rows=fvol.shape[0]+1, cols=fvol.shape[1]+1)
                table.style = 'Table_Style_NM'
                table_cells = table._cells
                for r in range (len(table.rows)):
                    row=table.rows[r]
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                    if r==0:
                        row.height = Cm(1.1)
                    else:
                        row.height = Cm(0.7)
                    
                table.cell(0,0).paragraphs[0].style="Table1_Header"
                table.cell(0,0).width = 1242000.2592
                table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
                table.columns[0].width= Cm(3.45)
                table.columns[-1].width=Cm(1.97)

                for k1 in range(1,len(f.columns)+1):
                    table.cell(0,k1).paragraphs[0].style="Table1_Header"
                    table.cell(0,k1).paragraphs[0].text=(fvol.columns[k1-1])
                for i in range(1,fvol.shape[0]+1):
                    table_cells[ i*(fvol.shape[1]+1)].paragraphs[0].style="Table1_Left"
                    table_cells[ i*(fvol.shape[1]+1)].paragraphs[0].text=(fvol.index[i-1])
                    for j in range(1,f.shape[1]+1):
                        table_cells[j + i * (fvol.shape[1]+1)].paragraphs[0].style="Table1_Right"
                        if (j %(f.shape[1]))==0:
                            table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                            table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2))+"%")
                            
                        elif (j %(f.shape[1]+1))==0:
                            table_cells[j + i * (fvol.shape[1]+1)].width = 1242000.2592
                            table_cells[j + i * (fvol.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2)))
                            
                        else:
                            table_cells[j + i * (fvol.shape[1]+1)].width = 594000.00072
                            table_cells[j + i * (fvol.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2)))

                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis",style="Source_NM")
                

                doc.add_paragraph(writeupvol(fvol,pdlist[k]["name"],market,volume),style="Normal")
              
                # doc.add_page_break()



                


            else:
                

                f['CAGR % '+"("+f.columns[0]+"-"+f.columns[-1]+")"] = (f.iloc[:, -1].div(f.iloc[:, 0]).pow(1./(len(f.columns) - 1)).sub(1))*100
                fvol['CAGR % '+"("+fvol.columns[1]+"-"+fvol.columns[-1]+")"] = (fvol.iloc[:, -1].div(fvol.iloc[:, 1]).pow(1./(len(fvol.columns) - 1)).sub(1))*100
                lisemp=[]
                for i in f.columns:
                    if i ==f.columns[0] or i == f.columns[-1]:
                        lisemp.append("")
                    else:
                        ar=((float(f.iloc[0][i])-float(f.iloc[0][str(int(i)-1)]))/float(f.iloc[0][str(int(i)-1)]))*100
                        lisemp.append(str(round(ar,2))+"%")

                f.loc["AGR%"]= lisemp

                lisemp=[]
                for i in fvol.columns:
                    if i ==fvol.columns[0] or i == fvol.columns[-1]:
                        lisemp.append("")
                    else:
                        ar=((float(fvol.iloc[0][i])-float(fvol.iloc[0][str(int(i)-1)]))/float(fvol.iloc[0][str(int(i)-1)]))*100
                        lisemp.append(str(round(ar,2))+"%")

                fvol.loc["AGR%"]= lisemp


                if pdlist[k]["heading"]==1:
                    h1=[None,None,str(count1)+". "+(k.split("#")[0]).upper()]
                    count1+=1
                    count2=1
                    head_row = format_row(h1, cellh1)
                    ws.append(head_row)
                    ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                    rnum +=1  

                elif pdlist[k]["heading"]==2 :
                    h2=[None,None,str(count2)+". "+(k.split("#")[0]).upper()]
                    count2+=1
                    head2_row = format_row(h2, cellh2)
                    ws.append(head2_row)
                    ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                    rnum +=1  

                            
                r1=[None,None,"TABLE "+str(count)+". "+("".join(mt)).upper()]
                count+=1
                row1=[None,None] 
                first_row = format_row(r1, cell)
                ws.append(first_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1  
                rowtable=rnum
                df=list(dataframe_to_rows(f, index=True, header=True))
                for i in range (len(df)):
                    row1=[None,None]
                    row= row1+df[i]
                #         print((row))
                    if i==0:
                        row=list(map(str, df[i]))
                        row[0]=pdlist[k]["comp"]
                        row=row1+row
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  
                    elif i==1:
                        continue
                    else:
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  

                tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
                tabname+=1
                headings = [pdlist[k]["comp"]]+list(f.columns)
                tab1._initialise_columns()
                for column, value1 in zip(tab1.tableColumns, headings):
                    column.name = value1
            #     tab1.headerRowCount = False
                
                tab1.tableStyleInfo = styletable
                ws.add_table(tab1)
                sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
                sourcerow = format_row(sourcer, cell1)
                ws.append(sourcerow)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                ws.append([None])
                ws.append([None])
                rnum +=3


                r1=[None,None,"TABLE "+str(count)+". "+("".join(mtv)).upper()]
                count+=1
                row1=[None,None] 
                first_row = format_row(r1, cell)
                ws.append(first_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1  
                rowtable=rnum
                df=list(dataframe_to_rows(fvol, index=True, header=True))
                for i in range (len(df)):
                    row1=[None,None]
                    row= row1+df[i]
                #         print((row))
                    if i==0:
                        row=list(map(str, df[i]))
                        row[0]=pdlist[k]["comp"]
                        row=row1+row
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  
                    elif i==1:
                        continue
                    else:
                        roe_num=format_row(row,cellnum)
                        ws.append(roe_num)
                        rnum +=1  

                tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
                tabname+=1
                headings = [pdlist[k]["comp"]]+list(f.columns)
                tab1._initialise_columns()
                for column, value1 in zip(tab1.tableColumns, headings):
                    column.name = value1
            #     tab1.headerRowCount = False
                
                tab1.tableStyleInfo = styletable
                ws.add_table(tab1)
                sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
                sourcerow = format_row(sourcer, cell1)
                ws.append(sourcerow)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                ws.append([None])
                ws.append([None])
                rnum +=3


                    

                if (tablesplit):
                    f=f.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)
                    fvol=fvol.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)


                f1=f.drop([f.columns[-1]],axis=1)
                f1=f1.drop(index=("AGR%"))
                p=doc.add_paragraph("".join(mt),style="FIG_no_NM")
                pic=plot(f1.T,t)
                ps=doc.add_picture(pic,width=Cm(17.2), height=Cm(8.51))
                os.remove(pic)
                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")
                # doc.add_paragraph(" ")
               

                p=doc.add_paragraph("".join(mt),style="Table_No_NM")
                table = doc.add_table(rows=f.shape[0]+1, cols=f.shape[1]+1)
                table.style = 'Table_Style_NM'
                table_cells = table._cells
                for r in range (len(table.rows)):
                    row=table.rows[r]
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                    if r==0:
                        row.height = Cm(1.1)
                    else:
                        row.height = Cm(0.7)

                table.cell(0,0).paragraphs[0].style="Table1_Header"
                table.cell(0,0).width = 1242000.2592
                table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
                table.columns[0].width= Cm(3.45)
                table.columns[-1].width=Cm(1.97)

                for k1 in range(1,len(f.columns)+1):
                    table.cell(0,k1).paragraphs[0].style="Table1_Header"
                    table.cell(0,k1).paragraphs[0].text=(f.columns[k1-1])
                for i in range(1,f.shape[0]+1):
                    table_cells[ i*(f.shape[1]+1)].paragraphs[0].style="Table1_Left"
                    table_cells[ i*(f.shape[1]+1)].paragraphs[0].text=(f.index[i-1])
                    for j in range(1,f.shape[1]+1):
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].style="Table1_Right"
                        if (j %(f.shape[1]))==0:
                            try:
                                table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                                table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2))+"%")

                            except:
                                table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                                table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(f.values[i-1][j-1]))


                        elif (j %(f.shape[1]+1))==0:
                            table_cells[j + i * (f.shape[1]+1)].width = 1242000.2592  
                            table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))
                                                
                        else:
                            try:
                                table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                                table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))
                                
                            except:
                                table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                                table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=str(f.values[i-1][j-1])


                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis",style="Source_NM")
                

                doc.add_paragraph(writeup(f,pdlist[k]["name"],market,value),style="Normal")
                
                


                


                f1vol=fvol.drop([fvol.columns[-1]],axis=1)
                f1vol=f1vol.drop(index=("AGR%"))
                p=doc.add_paragraph("".join(mtv),style="FIG_no_NM")
                picvol=plot(f1vol.T,t)
                ps=doc.add_picture(picvol,width=Cm(17.2), height=Cm(8.51))
                os.remove(picvol)
                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")
                # doc.add_paragraph(" ")

                
                p=doc.add_paragraph("".join(mtv),style="Table_No_NM")
                table = doc.add_table(rows=fvol.shape[0]+1, cols=fvol.shape[1]+1)
                table.style = 'Table_Style_NM'
                table_cells = table._cells
                for r in range (len(table.rows)):
                    row=table.rows[r]
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                    if r==0:
                        row.height = Cm(1.1)
                    else:
                        row.height = Cm(0.7)

                table.cell(0,0).paragraphs[0].style="Table1_Header"
                table.cell(0,0).width = 1242000.2592
                table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
                table.columns[0].width= Cm(3.45)
                table.columns[-1].width=Cm(1.97)

                for k1 in range(1,len(fvol.columns)+1):
                    table.cell(0,k1).paragraphs[0].style="Table1_Header"
                    table.cell(0,k1).paragraphs[0].text=(fvol.columns[k1-1])
                for i in range(1,fvol.shape[0]+1):
                    table_cells[ i*(fvol.shape[1]+1)].paragraphs[0].style="Table1_Left"
                    table_cells[ i*(fvol.shape[1]+1)].paragraphs[0].text=(fvol.index[i-1])
                    for j in range(1,fvol.shape[1]+1):
                        table_cells[j + i * (fvol.shape[1]+1)].paragraphs[0].style="Table1_Right"
                        if (j %(f.shape[1]))==0:
                            try:
                                table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                                table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2))+"%")

                            except:
                                table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                                table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(f.values[i-1][j-1]))
                            
                        elif (j %(fvol.shape[1]+1))==0:
                            table_cells[j + i * (fvol.shape[1]+1)].width = 1242000.2592  
                            table_cells[j + i * (fvol.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2)))
                                                
                        else:
                            try:
                                table_cells[j + i * (fvol.shape[1]+1)].width = 594000.00072
                                table_cells[j + i * (fvol.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2)))
                                
                            except:
                                table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                                table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=str(fvol.values[i-1][j-1])



                doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis",style="Source_NM")
                

                doc.add_paragraph(writeupvol(f,pdlist[k]["name"],market,volume),style="Normal")
                
            

        
        except Exception as e:
            dir= os.getcwd()
            logfile=dir+"\error_log.txt"
            file=open(logfile,"a")
            
            # data1={data}
            file.write("/n"+market.upper()+" "+reportname.upper()+" "+str(e))
            print(e)
            # f= open(".txt","a")

            continue    # doc.add_page_break()

    doc.save(dir+"/files/"+market+" "+reportname+" Value and Volume Market Report_Forecast Till "+str(g[1])+".docx")
    wb.save(dir+"/files/"+market+" "+reportname+"  Value and Volume Market Report_Forecast Excel doc Till "+str(g[1])+".xlsx")
    return (("/files/"+market+" "+reportname+" Value and Volume Market Report_Forecast Till "+str(g[1])+".docx","/files/"+market+" "+reportname+"  Value and Volume Market Report_Forecast Excel doc Till "+str(g[1])+".xlsx"))



def reportmake(toc1,year,ASP,reportname,market,value,volume,tablesplit):
    reportname=reportname.split("(1)")[0]
    v=list(year.keys())
    # yearvol={i:float(year[i])/float(ASP[i]) for i in year}
    for i in year:
        sk=year[i]
        year[i]=sk.replace(" ","")
        year[i]=sk.replace(",","")
    for n in ASP:
        sk=ASP[n]
        ASP[n]=sk.replace(" ","")
        ASP[n]=sk.replace(",","")
    if volume=="Units":
        yearvol={i:int(round(float(ASP[i]))) for i in ASP}
    else:
        yearvol=ASP
    g=[int(v[0]),int(v[-1])]
    toc1=tocl2vol(toc1,year,yearvol,g)
    # print(toc1)
    toc1=tocl3vol(toc1,g)
    toc1=tocl4vol(toc1,g)

    pdlist=abc(toc1,reportname,market)

    return(createreport(pdlist,reportname,market,value,volume,tablesplit,g))



def reportmakedownedit(toc1,market,reportname,value,volume,oggeo,regdet,tablesplit,startyear,endyear):
    reportname=reportname.split("(1)")[0]
    if oggeo=="Regional":
        for i in toc1:
            for j in i["tocreg"]:
                if j["country"]==market:
                    toc1=(j["toc"])
    else:
        for i in regdet:
            if market in regdet[i]:
                region=i
        
        for i in toc1:
            if i["region"]==region:
                for j in i["tocreg"]:
                    if j["country"]==market:
                        toc1=j["toc"]
                        break
            
    # print(toc1)
    g=[startyear,endyear]
    pdlist=abc(toc1,reportname,market)
    return(createreport(pdlist,reportname,market,value,volume,tablesplit,g))


def reportmakedownedityear(toc1,market,reportname,value,volume,oggeo,regdet,tablesplit,year,startyear,endyear):
    reportname=reportname.split("(1)")[0]
    if oggeo=="Regional":
        for i in toc1:
            if i["country"]==market:
                toc1=(i["toc"])
        tocspl(toc1,startyear,year)
    elif oggeo=="Country":
        tocspl(toc1,startyear,year)
    else:
        for i in regdet:
            for j in regdet[i]:
                if market == j["country"]:
                    region=i
        for i in toc1:
            if i["region"]==region:
                for j in i["tocreg"]:
                    if j["country"]==market:
                        toc1=j["toc"]
                        break
        tocspl(toc1,startyear,year)

    
    g=[year,endyear]
    pdlist=abc(toc1,reportname,market)
    return(createreport(pdlist,reportname,market,value,volume,tablesplit,g))