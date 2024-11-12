import pandas as pd
import json
from collections import Counter
import matplotlib
import matplotlib.pyplot as plt
from docx.shared import Inches,Cm
import docx
from docx.shared import RGBColor
import random
from docx.enum.table import WD_ROW_HEIGHT_RULE
import datetime
# import time
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, Border, Side, NamedStyle,Alignment
from openpyxl.cell.cell import WriteOnlyCell
import os
from openpyxl.worksheet.table import Table, TableStyleInfo
from mpmath import *

matplotlib.use('Agg')


def auto100(regen):
    regend=[]
    regen=[mpf(str(i).replace("%","")) for i in regen]
    sm=sum(regen)
    
    if sm!=100.0:
        if sm>100.0:
            diff=mpf((sm-100.0)/len(regen))
            
            for i in regen:
                val=i-diff
                if val >0:
                    regend.append(val)
                else:
                    regend.append(i)

        else :
            diff=mpf((100.0-sm)/len(regen))
            for i in regen:
                regend.append(i+diff)
    else:
        regend=regen
    return(regend)

def global100(toc):
    spstart=[]
    spend=[]
    spstartvol=[]
    spendvol=[]
    for i in toc:
        start=mpf((str(i["split"][0])).replace("%",""))
        end=mpf((str(i["split"][1])).replace("%",""))
        spstart.append(start)
        spend.append(end)
        startvol=mpf((str(i["vsplit"][0])).replace("%",""))
        endvol=mpf((str(i["vsplit"][1])).replace("%",""))
        spstartvol.append(startvol)
        spendvol.append(endvol)
    
    regst1=auto100(spstart)
    regen1=auto100(spend)
    regst1vol=auto100(spstartvol)
    regen1vol=auto100(spendvol)
    for i in range(len(toc)):
        toc[i]["split"][0]=regst1[i]
        toc[i]["split"][1]=regen1[i]
        toc[i]["vsplit"][0]=regst1vol[i]
        toc[i]["vsplit"][1]=regen1vol[i]
       

        
        

        
def region100(toc):
    for i in toc:
        spstart=[]
        spend=[]
        spstartvol=[]
        spendvol=[]
        for j in i["tocreg"]:
            start=mpf((str(j["split"][0])).replace("%",""))
            end=mpf((str(j["split"][1])).replace("%",""))
            spstart.append(start)
            spend.append(end)
            startvol=mpf((str(j["vsplit"][0])).replace("%",""))
            endvol=mpf((str(j["vsplit"][1])).replace("%",""))
            spstartvol.append(startvol)
            spendvol.append(endvol)
            
        regst1=auto100(spstart)
        regen1=auto100(spend)
        regst1vol=auto100(spstartvol)
        regen1vol=auto100(spendvol)
        for j in range(len(i["tocreg"])):
            i["tocreg"][j]["split"][0]=regst1[j]
            i["tocreg"][j]["split"][1]=regen1[j]
            i["tocreg"][j]["vsplit"][0]=regst1vol[j]
            i["tocreg"][j]["vsplit"][1]=regen1vol[j]


def toc100(toc):
    for i in toc:
        for j in i["tocreg"]:
            for k in j["toc"]: 
                sttc2=[]
                endtc2=[]
                sttc2vol=[]
                endtc2vol=[]
                for tc2 in k["toc_component"]:
                    sttc2.append(tc2["split"][0])
                    endtc2.append(tc2["split"][1])
                    sttc2vol.append(tc2["vsplit"][0])
                    endtc2vol.append(tc2["vsplit"][1])
                    if tc2["toc2_component"]!=[]:
                        sttc3=[]
                        endtc3=[]
                        sttc3vol=[]
                        endtc3vol=[]
                        for tc3 in tc2["toc2_component"]:
                            sttc3.append(tc3["split"][0])
                            endtc3.append(tc3["split"][1])
                            sttc3vol.append(tc3["vsplit"][0])
                            endtc3vol.append(tc3["vsplit"][1])
                            if tc3["toc3_component"]!=[]:
                                sttc4=[]
                                endtc4=[]
                                sttc4vol=[]
                                endtc4vol=[]
                                for tc4 in tc3["toc3_component"]:
                                    # print(tc4)
                                    sttc4.append(tc4["split"][0])
                                    endtc4.append(tc4["split"][1])
                                    sttc4vol.append(tc4["vsplit"][0])
                                    endtc4vol.append(tc4["vsplit"][1])
                                
                                
                                regst13=auto100(sttc4)
                                regen13=auto100(endtc4)
                                regst13vol=auto100(sttc4vol)
                                regen13vol=auto100(endtc4vol)
                                for tc4 in range(len(tc3["toc3_component"])):
                                    tc3["toc3_component"][tc4]["split"][0]=regst13[tc4]
                                    tc3["toc3_component"][tc4]["split"][1]=regen13[tc4]
                                    tc3["toc3_component"][tc4]["vsplit"][0]=regst13vol[tc4]
                                    tc3["toc3_component"][tc4]["vsplit"][1]=regen13vol[tc4]
                               

                        regst12=auto100(sttc3)
                        regen12=auto100(endtc3)
                        regst12vol=auto100(sttc3vol)
                        regen12vol=auto100(endtc3vol)
                        for tc3 in range(len(tc2["toc2_component"])):
                            tc2["toc2_component"][tc3]["split"][0]=regst12[tc3]
                            tc2["toc2_component"][tc3]["split"][1]=regen12[tc3]
                            tc2["toc2_component"][tc3]["vsplit"][0]=regst12vol[tc3]
                            tc2["toc2_component"][tc3]["vsplit"][1]=regen12vol[tc3]


                regst1=auto100(sttc2)
                regen1=auto100(endtc2) 
                regst1vol=auto100(sttc2vol)
                regen1vol=auto100(endtc2vol) 
                for tc2 in range(len(k["toc_component"])):
                    k["toc_component"][tc2]["split"][0]=regst1[tc2]
                    k["toc_component"][tc2]["split"][1]=regen1[tc2]
                    k["toc_component"][tc2]["vsplit"][0]=regst1vol[tc2]
                    k["toc_component"][tc2]["vsplit"][1]=regen1vol[tc2]



                

            
        
    
    

    

def global1(toc,year,yearvol,g):
    for i in toc:
        s1=mpf((str(i["split"][0])).replace("%",""))
        s2=mpf((str(i["split"][1])).replace("%",""))
        s1v=mpf((str(i["vsplit"][0])).replace("%",""))
        s2v=mpf((str(i["vsplit"][1])).replace("%",""))
        step=(s2-s1)/(g[1]-g[0])
        stepv=(s2v-s1v)/(g[1]-g[0])
        i["split"]={}
        i["vsplit"]={}
        for k in range (g[0],g[1]+1): 
            i["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
            i["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
            # float((s1*mpf(yearvol[str(k)]))/100)
            s1+=step
            s1v+=stepv
        region(i['tocreg'],i["split"],i["vsplit"],g)

def region(toc,year,yearvol,g):
    for i in toc:
        s1=mpf((str(i["split"][0])).replace("%",""))
        s2=mpf((str(i["split"][1])).replace("%",""))
        s1v=mpf((str(i["vsplit"][0])).replace("%",""))
        s2v=mpf((str(i["vsplit"][1])).replace("%",""))
        step=(s2-s1)/(g[1]-g[0])
        stepv=(s2v-s1v)/(g[1]-g[0])
        i["split"]={}
        i["vsplit"]={}
        for k in range (g[0],g[1]+1):
            i["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
            i["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
            s1+=step
            s1v+=stepv
        tocl2reg(i["toc"],i["split"],i["vsplit"],g)
        tocl3reg(i["toc"],g)
        tocl4reg(i["toc"],g)
        


def tocl2reg(toc,year1,yearvol,g):
    for i in toc:
        for j in (i["toc_component"]):
            s1=mpf((str(j["split"][0])).replace("%",""))
            s2=mpf((str(j["split"][1])).replace("%",""))
            step=(s2-s1)/(g[1]-g[0])
            s1v=mpf((str(j["vsplit"][0])).replace("%",""))
            s2v=mpf((str(j["vsplit"][1])).replace("%",""))
            stepv=(s2v-s1v)/(g[1]-g[0])
            j["split"]={}
            j["vsplit"]={}
            for k in range (g[0],g[1]+1):
                j["split"][str(k)]=float((s1*mpf(str(year1[str(k)]).replace(",","")))/100)
                j["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
                s1+=step
                s1v+=stepv

def tocl3reg(toc,g):
    for j in range(len(toc)):
        for c in toc[j]["toc_component"]:
            year=c["split"]
            yearvol=c["vsplit"]
            if (len(c["toc2_component"]))!=0:
                for t in (c["toc2_component"]):
                    s1=mpf((str(t["split"][0])).replace("%",""))
                    s2=mpf((str(t["split"][1])).replace("%",""))
                    step=(s2-s1)/(g[1]-g[0])
                    s1v=mpf((str(t["vsplit"][0])).replace("%",""))
                    s2v=mpf((str(t["vsplit"][1])).replace("%",""))
                    stepv=(s2v-s1v)/(g[1]-g[0])
                    t["split"]={}
                    t["vsplit"]={}
                    for k in range (g[0],(g[1]+1)):
                        t["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
                        t["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
                        s1+=step
                        s1v+=stepv
                        
                    
def tocl4reg(toc,g):
    for j in range(len(toc)):
        for c in toc[j]["toc_component"]:
            for t in (c["toc2_component"]):
                year=t["split"]
                yearvol=t["vsplit"]
                if (len(t["toc3_component"]))!=0:
                    for r in (t["toc3_component"]):
                        s1=mpf((str(r["split"][0])).replace("%",""))
                        s2=mpf((str(r["split"][1])).replace("%",""))
                        step=(s2-s1)/(g[1]-g[0])
                        s1v=mpf((str(r["vsplit"][0])).replace("%",""))
                        s2v=mpf((str(r["vsplit"][1])).replace("%",""))
                        stepv=(s2v-s1v)/(g[1]-g[0])
                        r['split']={}
                        r['vsplit']={}
                        for k in range (g[0],(g[1]+1)):
                            r['split'][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
                            r['vsplit'][str(k)]=float((s1*mpf(str(yearvol[str(k)]).replace(",","")))/100)
                            s1+=step 
                            s1v+=stepv 



def globalfrml2(toc,l2):
    k1=Counter()
    for k in toc:
        for v in k["tocreg"]:
            for z in (v["toc"]):
                for t in z["toc_component"]:
                    if (t["toc2_name"]) ==l2:
                        k1+=(Counter(t["split"]))
    return((k1))
def l2byreg(toc,reg,l1,l2):
    k1=Counter()
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                for j in v["toc"]:
                    if j["toc_name"]==l1:
                        for c in j["toc_component"]:
                            if c["toc2_name"]==l2:
                                k1+=(Counter(c["split"]))
    return(k1)

def l2bycountry(toc,cun,reg,l1,l2):
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                if v["country"]==cun:
                    for j in v["toc"]:
                        if j["toc_name"]==l1:
                            for c in j["toc_component"]:
                                if c["toc2_name"]==l2:
                                    return(c["split"])
def globalfrml3(toc,l2,l3):
    k1=Counter()
    for k in toc:
        for v in k["tocreg"]:
            for z in (v["toc"]):
                for t in z["toc_component"]:
                    if (t["toc2_name"]) ==l2:
                        for c in t["toc2_component"]:
                            if c["toc3_name"]==l3:
                                k1+=(Counter(c["split"]))
    return((k1))


def l3byreg(toc,reg,l1,l2,l3):
    k1=Counter()
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                for j in v["toc"]:
                    if j["toc_name"]==l1:
                        for c in j["toc_component"]:
                            if c["toc2_name"]==l2:
                                for c1 in c["toc2_component"]:
                                    if c1["toc3_name"]==l3:
                                        k1+=(Counter(c1["split"]))
    return(k1)

def l3bycountry(toc,cun,reg,l1,l2,l3):
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                if v["country"]==cun:
                    for j in v["toc"]:
                        if j["toc_name"]==l1:
                            for c in j["toc_component"]:
                                if c["toc2_name"]==l2:
                                    for c1 in c["toc2_component"]:
                                        if c1["toc3_name"]==l3:
                                            return(c1["split"])
                                        
                                        
                                        
def globalfrml4(toc,l2,l3,l4):
    k1=Counter()
    for k in toc:
        for v in k["tocreg"]:
            for z in (v["toc"]):
                for t in z["toc_component"]:
                    if (t["toc2_name"]) ==l2:
                        for c in t["toc2_component"]:
                            if c["toc3_name"]==l3:
                                for c1 in c["toc3_component"]:
                                    if c1["toc4_name"]==l4:
                                        k1+=(Counter(c1["split"]))
    return((k1))


def l4byreg(toc,reg,l1,l2,l3,l4):
    k1=Counter()
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                for j in v["toc"]:
                    if j["toc_name"]==l1:
                        for c in j["toc_component"]:
                            if c["toc2_name"]==l2:
                                for c1 in c["toc2_component"]:
                                    if c1["toc3_name"]==l3:
                                        for n in c1["toc3_component"]:
                                             if n["toc4_name"]==l4:
                                                k1+=(Counter(n["split"]))
    return(k1)

def l4bycountry(toc,cun,reg,l1,l2,l3,l4):
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                if v["country"]==cun:
                    for j in v["toc"]:
                        if j["toc_name"]==l1:
                            for c in j["toc_component"]:
                                if c["toc2_name"]==l2:
                                    for c1 in c["toc2_component"]:
                                        if c1["toc3_name"]==l3:
                                            for n in c1["toc3_component"]:
                                                if n["toc4_name"]==l4:
                                                    return(n["split"])

def globalbyregion(toc,pdlist,market):
    pdlist["global "+market+" by region"]={"heading":1,"before":"","sin":0,"comp":"Region","data":{},"name":"Global"}
    for reg1 in toc: 
        pdlist["global "+market+" by region"]["data"][reg1["region"]]=reg1["split"]
    

def regionbytocl2(re,l1,l2,toc):
    k1=Counter()
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                for k in j["toc"]:
                    if k["toc_name"]==l1:
                        for n in k["toc_component"]:
                            if  n["toc2_name"]==l2:
                                k1+=n["split"]
    return(k1)

def regionbytocl3(re,l1,l2,l3,toc):
    k1=Counter()
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                for k in j["toc"]:
                    if k["toc_name"]==l1:
                        for n in k["toc_component"]:
                            if  n["toc2_name"]==l2:
                                for n1 in n["toc2_component"]:
                                    if n1["toc3_name"]==l3:
                                        k1+=n1["split"]
    return(k1)

def regionbytocl4(re,l1,l2,l3,l4,toc):
    k1=Counter()
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                for k in j["toc"]:
                    if k["toc_name"]==l1:
                        for n in k["toc_component"]:
                            if  n["toc2_name"]==l2:
                                for n1 in n["toc2_component"]:
                                    if n1["toc3_name"]==l3:
                                        for n2 in n1["toc3_component"]:
                                            if n2["toc4_name"]==l4:
                                                k1+=n2["split"]
    return(k1)

def regcountsplit(reg,toc,country):
    for l1 in toc:
        if l1["region"]==reg:
            for l2 in l1["tocreg"]:
                if l2["country"]==country:
                    return(l2["split"])
            


def globalfrml2vol(toc,l2):
    k1=Counter()
    for k in toc:
        for v in k["tocreg"]:
            for z in (v["toc"]):
                for t in z["toc_component"]:
                    if (t["toc2_name"]) ==l2:
                        k1+=(Counter(t["vsplit"]))
    return((k1))
def l2byregvol(toc,reg,l1,l2):
    k1=Counter()
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                for j in v["toc"]:
                    if j["toc_name"]==l1:
                        for c in j["toc_component"]:
                            if c["toc2_name"]==l2:
                                k1+=(Counter(c["vsplit"]))
    return(k1)

def l2bycountryvol(toc,cun,reg,l1,l2):
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                if v["country"]==cun:
                    for j in v["toc"]:
                        if j["toc_name"]==l1:
                            for c in j["toc_component"]:
                                if c["toc2_name"]==l2:
                                    return(c["vsplit"])
def globalfrml3vol(toc,l2,l3):
    k1=Counter()
    for k in toc:
        for v in k["tocreg"]:
            for z in (v["toc"]):
                for t in z["toc_component"]:
                    if (t["toc2_name"]) ==l2:
                        for c in t["toc2_component"]:
                            if c["toc3_name"]==l3:
                                k1+=(Counter(c["vsplit"]))
    return((k1))


def l3byregvol(toc,reg,l1,l2,l3):
    k1=Counter()
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                for j in v["toc"]:
                    if j["toc_name"]==l1:
                        for c in j["toc_component"]:
                            if c["toc2_name"]==l2:
                                for c1 in c["toc2_component"]:
                                    if c1["toc3_name"]==l3:
                                        k1+=(Counter(c1["vsplit"]))
    return(k1)

def l3bycountryvol(toc,cun,reg,l1,l2,l3):
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                if v["country"]==cun:
                    for j in v["toc"]:
                        if j["toc_name"]==l1:
                            for c in j["toc_component"]:
                                if c["toc2_name"]==l2:
                                    for c1 in c["toc2_component"]:
                                        if c1["toc3_name"]==l3:
                                            return(c1["vsplit"])
                                        
                                        
                                        
def globalfrml4vol(toc,l2,l3,l4):
    k1=Counter()
    for k in toc:
        for v in k["tocreg"]:
            for z in (v["toc"]):
                for t in z["toc_component"]:
                    if (t["toc2_name"]) ==l2:
                        for c in t["toc2_component"]:
                            if c["toc3_name"]==l3:
                                for c1 in c["toc3_component"]:
                                    if c1["toc4_name"]==l4:
                                        k1+=(Counter(c1["vsplit"]))
    return((k1))


def l4byregvol(toc,reg,l1,l2,l3,l4):
    k1=Counter()
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                for j in v["toc"]:
                    if j["toc_name"]==l1:
                        for c in j["toc_component"]:
                            if c["toc2_name"]==l2:
                                for c1 in c["toc2_component"]:
                                    if c1["toc3_name"]==l3:
                                        for n in c1["toc3_component"]:
                                             if n["toc4_name"]==l4:
                                                k1+=(Counter(n["vsplit"]))
    return(k1)

def l4bycountryvol(toc,cun,reg,l1,l2,l3,l4):
    for k in toc:
        if k["region"]==reg:
            for v in k["tocreg"]:
                if v["country"]==cun:
                    for j in v["toc"]:
                        if j["toc_name"]==l1:
                            for c in j["toc_component"]:
                                if c["toc2_name"]==l2:
                                    for c1 in c["toc2_component"]:
                                        if c1["toc3_name"]==l3:
                                            for n in c1["toc3_component"]:
                                                if n["toc4_name"]==l4:
                                                    return(n["vsplit"])


    

def regionbytocl2vol(re,l1,l2,toc):
    k1=Counter()
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                for k in j["toc"]:
                    if k["toc_name"]==l1:
                        for n in k["toc_component"]:
                            if  n["toc2_name"]==l2:
                                k1+=n["vsplit"]
    return(k1)

def regionbytocl3vol(re,l1,l2,l3,toc):
    k1=Counter()
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                for k in j["toc"]:
                    if k["toc_name"]==l1:
                        for n in k["toc_component"]:
                            if  n["toc2_name"]==l2:
                                for n1 in n["toc2_component"]:
                                    if n1["toc3_name"]==l3:
                                        k1+=n1["vsplit"]
    return(k1)

def regionbytocl4vol(re,l1,l2,l3,l4,toc):
    k1=Counter()
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                for k in j["toc"]:
                    if k["toc_name"]==l1:
                        for n in k["toc_component"]:
                            if  n["toc2_name"]==l2:
                                for n1 in n["toc2_component"]:
                                    if n1["toc3_name"]==l3:
                                        for n2 in n1["toc3_component"]:
                                            if n2["toc4_name"]==l4:
                                                k1+=n2["vsplit"]
    return(k1)

def regcountsplitvol(reg,toc,country):
    for l1 in toc:
        if l1["region"]==reg:
            for l2 in l1["tocreg"]:
                if l2["country"]==country:
                    return(l2["vsplit"])
            




def countrybytocl2(re,coun,l1,l2,toc):
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                if j["country"]==coun:
                    for k in j["toc"]:
                        if k["toc_name"]==l1:
                            for n in k["toc_component"]:
                                if  n["toc2_name"]==l2:
                                    return(n["split"])
                                
def countrybytocl3(re,coun,l1,l2,l3,toc):
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                if j["country"]==coun:
                    for k in j["toc"]:
                        if k["toc_name"]==l1:
                            for n in k["toc_component"]:
                                if  n["toc2_name"]==l2:
                                    for n1 in n["toc2_component"]:
                                        if  n1["toc3_name"]==l3:
                                            return(n1["split"])
                                        
def countrybytocl4(re,coun,l1,l2,l3,l4,toc):
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                if j["country"]==coun:
                    for k in j["toc"]:
                        if k["toc_name"]==l1:
                            for n in k["toc_component"]:
                                if  n["toc2_name"]==l2:
                                    for n1 in n["toc2_component"]:
                                        if  n1["toc3_name"]==l3:
                                            for n2 in n1["toc3_component"]:
                                                if  n2["toc4_name"]==l4:
                                                    return(n2["split"])
                                                
                    



def countrybytocl2vol(re,coun,l1,l2,toc):
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                if j["country"]==coun:
                    for k in j["toc"]:
                        if k["toc_name"]==l1:
                            for n in k["toc_component"]:
                                if  n["toc2_name"]==l2:
                                    return(n["vsplit"])
                                
def countrybytocl3vol(re,coun,l1,l2,l3,toc):
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                if j["country"]==coun:
                    for k in j["toc"]:
                        if k["toc_name"]==l1:
                            for n in k["toc_component"]:
                                if  n["toc2_name"]==l2:
                                    for n1 in n["toc2_component"]:
                                        if  n1["toc3_name"]==l3:
                                            return(n1["vsplit"])
                                        
def countrybytocl4vol(re,coun,l1,l2,l3,l4,toc):
    for i in toc:
        if i["region"]==re:
            for j in i["tocreg"]:
                if j["country"]==coun:
                    for k in j["toc"]:
                        if k["toc_name"]==l1:
                            for n in k["toc_component"]:
                                if  n["toc2_name"]==l2:
                                    for n1 in n["toc2_component"]:
                                        if  n1["toc3_name"]==l3:
                                            for n2 in n1["toc3_component"]:
                                                if  n2["toc4_name"]==l4:
                                                    return(n2["vsplit"])
                                                
                    






def globalbyregion(toc,pdlist,market,count):
    pdlist["GLOBAL "+market+" Market, BY REGION"+" #"+str(count)]={"heading":1,"before":"","sin":0,"comp":"Region","vdata":{},"data":{},"name":"Global "+market,"val":"GLOBAL "+market+" Market VALUE, BY REGION","vol":"GLOBAL "+market+" Market VOLUME, BY REGION"}
    for reg1 in toc: 
        pdlist["GLOBAL "+market+" Market, BY REGION"+" #"+str(count)]["data"][reg1["region"]]=reg1["split"]
        pdlist["GLOBAL "+market+" Market, BY REGION"+" #"+str(count)]["vdata"][reg1["region"]]=reg1["vsplit"]
    count+=1
    return count
def addl1byglob(l1,tocl1,pdlist,market,toc,count):
    pdlist["Global "+market+" Market, By "+l1+" #"+str(count)]={"heading":1,"before":"","sin":0,"comp":l1,"data":{},"vdata":{},"name":"Global "+market,"val":"Global "+market+" Market VALUE, By "+l1,"vol":"Global "+market+" Market VOLUME, By "+l1}
    for l2 in (tocl1[l1].keys()):
        pdlist["Global "+market+" Market, By "+l1+" #"+str(count)]["data"][l2]=globalfrml2(toc,l2)
        pdlist["Global "+market+" Market, By "+l1+" #"+str(count)]["vdata"][l2]=globalfrml2vol(toc,l2)
    count+=1
    return count
def add2docl2(l1,l2,tocl1,reg,a,pdlist,toc,count):
    if a==0:
        pdlist[l2+" MARKET, BY REGION"+" #"+str(count)]={"heading":3,"before":l2,"sin":0,"comp":"Region","data":{},"name":l2,"vdata":{},"name":l2,"val":l2+" Market VALUE, BY REGION","vol":l2+" Market VOLUME, BY REGION"}
    else:
        pdlist[l2+" MARKET, BY REGION"+" #"+str(count)]={"heading":3,"before":"","sin":0,"comp":"Region","data":{},"name":l2,"vdata":{},"name":l2,"val":l2+" Market VALUE, BY REGION","vol":l2+" Market VOLUME, BY REGION"}
    for re in reg.keys():
        pdlist[l2+" MARKET, BY REGION"+" #"+str(count)]["data"][re]=l2byreg(toc,re,l1,l2)
        pdlist[l2+" MARKET, BY REGION"+" #"+str(count)]["vdata"][re]=l2byregvol(toc,re,l1,l2)
    count+=1
    # for m in reg:
    #     pdlist[m+" "+l2+" MARKET, By Country"+" #"+str(count)]={"heading":4,"before":"","sin":0,"comp":"Country","data":{},"vdata":{},"name":m+" "+l2,"val":m+" "+l2+" MARKET VALUE, By Country","vol":m+" "+l2+" MARKET VOLUME, By Country"}
    #     f=reg[m]
    #     for n in f:
    #         pdlist[m+" "+l2+" MARKET, By Country"+" #"+str(count)]["data"][n]=l2bycountry(toc,n,m,l1,l2)
    #         pdlist[m+" "+l2+" MARKET, By Country"+" #"+str(count)]["vdata"][n]=l2bycountryvol(toc,n,m,l1,l2)
    #     count+=1
    return count
            
def adddocl3global(l1,l2,tocl1,pdlist,toc,count):
    pdlist["GLOBAL "+l2+" MARKET, BY "+l1+" #"+str(count)]={"heading":3,"before":l2,"sin":0,"comp":l2,"data":{},"vdata":{},"name":"Global "+l2,"val":"GLOBAL "+l2+" MARKET VALUE, BY "+l1,"vol":"GLOBAL "+l2+" MARKET VOLUME, BY "+l1}
    for l3 in (tocl1[l1][l2].keys()):
        pdlist["GLOBAL "+l2+" MARKET, BY "+l1+" #"+str(count)]["data"][l3]=globalfrml3(toc,l2,l3)
        pdlist["GLOBAL "+l2+" MARKET, BY "+l1+" #"+str(count)]["vdata"][l3]=globalfrml3vol(toc,l2,l3)
    count+=1
    return count

def add2docl3(l1,l2,l3,tocl1,reg,pdlist,toc,count):
    pdlist[l3+" Market, By Region"+" #"+str(count)]={"heading":4,"before":"","sin":0,"comp":"Region","data":{},"vdata":{},"name":l3,"val":l3+" Market VALUE, By Region","vol":l3+" Market VOLUME, By Region"}
    for re in reg.keys():
        pdlist[l3+" Market, By Region"+" #"+str(count)]["data"][re]=l3byreg(toc,re,l1,l2,l3)
        pdlist[l3+" Market, By Region"+" #"+str(count)]["vdata"][re]=l3byregvol(toc,re,l1,l2,l3)
    count+=1

    # for re in reg:
    #     pdlist[re+" "+l3+" Market, By Country"+" #"+str(count)]={"heading":5,"sin":0,"comp":"Country","before":"","data":{},"vdata":{},"name":re+" "+l3,"val":re+" "+l3+" Market VALUE, By Country","vol":re+" "+l3+" Market VOLUME, By Country"} 
    #     f=reg[re]
    #     for n in f:
    #         pdlist[re+" "+l3+" Market, By Country"+" #"+str(count)]["data"][n]=l3bycountry(toc,n,re,l1,l2,l3)
    #         pdlist[re+" "+l3+" Market, By Country"+" #"+str(count)]["vdata"][n]=l3bycountryvol(toc,n,re,l1,l2,l3)
    #     count+=1
    return count
            

def adddocl4global(l1,l2,l3,tocl1,pdlist,toc,count):
    pdlist["GLOBAL "+l3+" MARKET, BY "+l2+" #"+str(count)]={"heading":5,"before":l3,"sin":0,"comp":l3,"data":{},"vdata":{},"name":"Global "+l3,"val":"GLOBAL "+l3+" MARKET VALUE, BY "+l2,"vol":"GLOBAL "+l3+" MARKET VOLUME, BY "+l2}
    for l4 in (tocl1[l1][l2][l3]):
        pdlist["GLOBAL "+l3+" MARKET, BY "+l2+" #"+str(count)]["data"][l4]=globalfrml4(toc,l2,l3,l4)
        pdlist["GLOBAL "+l3+" MARKET, BY "+l2+" #"+str(count)]["vdata"][l4]=globalfrml4vol(toc,l2,l3,l4)
    count+=1
    return count
    
def add2docl4(l1,l2,l3,l4,tocl1,reg,pdlist,toc,count):
    pdlist[l4+" Market, By Region"+" #"+str(count)]={"heading":6,"sin":0,"before":"","comp":"Region","data":{},"vdata":{},"name":l4,"val":l4+" Market VALUE, By Region","vol":l4+" Market VOLUME, By Region"}
    for re in reg.keys():
        pdlist[l4+" Market, By Region"+" #"+str(count)]["data"][re]=l4byreg(toc,re,l1,l2,l3,l4)
        pdlist[l4+" Market, By Region"+" #"+str(count)]["vdata"][re]=l4byregvol(toc,re,l1,l2,l3,l4)
    count+=1
        

    # for re in reg:
    #     pdlist[re+" "+l4+" Market, By Country"+" #"+str(count)]={"heading":7,"before":"","sin":0,"comp":"Country","data":{},"vdata":{},"name":re+" "+l4,"val":re+" "+l4+" Market VALUE, By Country","vol":re+" "+l4+" Market VOLUME, By Country"}
    #     f=reg[re]
    #     for n in f:
    #         pdlist[re+" "+l4+" Market, By Country"+" #"+str(count)]["data"][n]=l4bycountry(toc,n,re,l1,l2,l3,l4)
    #         pdlist[re+" "+l4+" Market, By Country"+" #"+str(count)]["vdata"][n]=l4bycountryvol(toc,n,re,l1,l2,l3,l4)
    #     count+=1
        
    return count
#         print(pdlist[re+" "+l4+" Market, By Country"])


def l1byreg(l1,tocl1,re,a,pdlist,market,toc,count):
    count1=count
    if a==0:
        pdlist[re+" "+market+" Market, By "+l1+" #"+str(count1)]={"heading":3,"before":re,"sin":0,"comp":l1,"data":{},"vdata":{},"name":re,"val":re+" "+market+" Market VALUE, By "+l1,"vol":re+" "+market+" Market VOLUME, By "+l1}
    else:
        pdlist[re+" "+market+" Market, By "+l1+" #"+str(count1)]={"heading":3,"before":"","sin":0,"comp":l1,"data":{},"vdata":{},"name":re,"val":re+" "+market+" Market VALUE, By "+l1,"vol":re+" "+market+" Market VOLUME, By "+l1}
    for l2 in (tocl1[l1]):
        pdlist[re+" "+market+" Market, By "+l1+" #"+str(count1)]["data"][l2]=regionbytocl2(re,l1,l2,toc)
        pdlist[re+" "+market+" Market, By "+l1+" #"+str(count1)]["vdata"][l2]=regionbytocl2vol(re,l1,l2,toc)
        if (tocl1[l1][l2])!={}:
            count2=count1
            pdlist[re+" "+l2+" Market, By "+l1+" #"+str(count2)]={"heading":4,"before":"","sin":0,"comp":l2,"data":{},"name":re+" "+l2,"vdata":{},"val":re+" "+l2+" Market VALUE, By "+l1,"vol":re+" "+l2+" Market VOLUME, By "+l1}
            for l3 in (tocl1[l1][l2]):
                pdlist[re+" "+l2+" Market, By "+l1+" #"+str(count2)]["data"][l3]=regionbytocl3(re,l1,l2,l3,toc)
                pdlist[re+" "+l2+" Market, By "+l1+" #"+str(count2)]["vdata"][l3]=regionbytocl3vol(re,l1,l2,l3,toc)
                if (tocl1[l1][l2][l3])!=[]:
                    count3=count2
                    pdlist[re+" "+l3+" Market, By "+l2+" #"+str(count3)]={"heading":5,"before":"","sin":0,"comp":l3,"data":{},"vdata":{},"name":re+" "+l3,"val":re+" "+l3+" Market VALUE, By "+l2,"vol":re+" "+l3+" Market VOLUME, By "+l2}
                    for l4 in (tocl1[l1][l2][l3]):
                        pdlist[re+" "+l3+" Market, By "+l2+" #"+str(count3)]["data"][l4]=regionbytocl4(re,l1,l2,l3,l4,toc)
                        pdlist[re+" "+l3+" Market, By "+l2+" #"+str(count3)]["vdata"][l4]=regionbytocl4vol(re,l1,l2,l3,l4,toc)
                    count3+=1
            count2+=1
    return count
    
                        
                        
                        
def regionbycountry(re,pdlist,market,toc,reg,count):
    pdlist[re+" "+market+" Market, By Country"+" #"+str(count)]={"heading":3,"before":"","sin":0,"comp":"Country","data":{},"name":re+" "+market,"vdata":{},"val":re+" "+market+" VALUE, By Country","vol":re+" "+market+" VOLUME, By Country"}
    for countr in reg[re]:
        pdlist[re+" "+market+" Market, By Country"+" #"+str(count)]["data"][countr]=regcountsplit(re,toc,countr)
        pdlist[re+" "+market+" Market, By Country"+" #"+str(count)]["vdata"][countr]=regcountsplitvol(re,toc,countr)
    return count+1
        
        
def l1bycountry(l1,tocl1,re,countr,a,pdlist,market,toc,count):
#     count1=count
    if a==0:
        pdlist[countr+" "+market+" Market, By "+l1+" #"+str(count)]={"heading":5,"before":countr,"sin":0,"comp":l1,"data":{},"vdata":{},"name":countr+" "+market,"val":countr+" "+market+" Market VALUE, By "+l1,"vol":countr+" "+market+" Market VOLUME, By "+l1}
    else:
        pdlist[countr+" "+market+" Market, By "+l1+" #"+str(count)]={"heading":5,"before":"","sin":0,"comp":l1,"data":{},"vdata":{},"name":countr+" "+market,"val":countr+" "+market+" Market VALUE, By "+l1,"vol":countr+" "+market+" Market VOLUME, By "+l1}
    for l2 in (tocl1[l1]):
        pdlist[countr+" "+market+" Market, By "+l1+" #"+str(count)]["data"][l2]=countrybytocl2(re,countr,l1,l2,toc)
        pdlist[countr+" "+market+" Market, By "+l1+" #"+str(count)]["vdata"][l2]=countrybytocl2vol(re,countr,l1,l2,toc)
        count2=count+1
        if (tocl1[l1][l2])!={}:
            
            pdlist[countr+" "+l2+" Market, By "+l1+" #"+str(count2)]={"heading":6,"before":"","comp":l2,"data":{},"vdata":{},"name":countr+" "+l2,"val":countr+" "+l2+" Market VALUE, By "+l1,"vol":countr+" "+l2+" Market VOLUME, By "+l1}
            for l3 in (tocl1[l1][l2]):
                pdlist[countr+" "+l2+" Market, By "+l1+" #"+str(count2)]["data"][l3]=countrybytocl3(re,countr,l1,l2,l3,toc)
                pdlist[countr+" "+l2+" Market, By "+l1+" #"+str(count2)]["vdata"][l3]=countrybytocl3vol(re,countr,l1,l2,l3,toc)
                count3=count2
                if (tocl1[l1][l2][l3])!=[]:
                   
                    pdlist[countr+" "+l3+" Market, By "+l2+" #"+str(count3)]={"heading":7,"before":"","sin":0,"comp":l3,"data":{},"vdata":{},"name":countr+" "+l3,"val":countr+" "+l3+" Market VALUE, By "+l2,"vol":countr+" "+l3+" Market VOLUME, By "+l2}
                    for l4 in (tocl1[l1][l2][l3]):
                        pdlist[countr+" "+l3+" Market, By "+l2+" #"+str(count3)]["data"][l4]=countrybytocl4(re,countr,l1,l2,l3,l4,toc)
                        pdlist[countr+" "+l3+" Market, By "+l2+" #"+str(count3)]["vdata"][l4]=countrybytocl4vol(re,countr,l1,l2,l3,l4,toc)
                    count3+=1
        else:
            count3=0
        count2+=1
    count=count+count2+count3
    return(count)

def makepdlist(toc,tocl1,reg,market):
    pdlist={}
    count=0
    for l1 in tocl1.keys():
        count=addl1byglob(l1,tocl1,pdlist,market,toc,count)
        for l2 in (tocl1[l1].keys()):
            if tocl1[l1][l2]=={}:
                # print(l2)
                count=add2docl2(l1,l2,tocl1,reg,0,pdlist,toc,count)
            else:
                count=adddocl3global(l1,l2,tocl1,pdlist,toc,count)
                for l3 in tocl1[l1][l2]:
                    
                    if tocl1[l1][l2][l3]!=[]:
                        
                        count=adddocl4global(l1,l2,l3,tocl1,pdlist,toc,count)
                        for l4 in tocl1[l1][l2][l3]:
                            
                            count=add2docl4(l1,l2,l3,l4,tocl1,reg,pdlist,toc,count)
                            
                    # print(l3)
                    count=add2docl3(l1,l2,l3,tocl1,reg,pdlist,toc,count)
                count=add2docl2(l1,l2,tocl1,reg,1,pdlist,toc,count)


    count=globalbyregion(toc,pdlist,market,count)
    for re in reg:
        a=0
        for l1 in tocl1:
            count=l1byreg(l1,tocl1,re,a,pdlist,market,toc,count)
            a=1
        count=regionbycountry(re,pdlist,market,toc,reg,count)
        for countr in reg[re]:
            b=0
            for l1 in tocl1:
                count=l1bycountry(l1,tocl1,re,countr,b,pdlist,market,toc,count)
                b=1

    return(pdlist)




def plot(j,t):
    dir=os.getcwd()
    ax=j.plot(kind="bar",figsize=(8,6.25))
    ax.set_axisbelow(True)
    ax.yaxis.grid(color='gray', linestyle='dashed')
    plt.xticks(rotation='horizontal')
    # if (j.shape[1]>=5):
    plt.legend(loc="upper left")
    #     # ,bbox_to_anchor=(0.5, -0.30),fancybox=True, shadow=True,handletextpad=0.5)
    # else:
    #     plt.legend(loc="upper left")
        # ,bbox_to_anchor=(0.5, -0.175),fancybox=True, shadow=True,handletextpad=0.5)
    plt.savefig(dir+"/temp/"+t+'example.png',bbox_inches='tight',pad_inches =0.05)
    plt.close()
    return(dir+"/temp/"+t+'example.png')




def parastyle(headno):
    if headno==1:
        return "Main_Heading-NM"
    elif headno==2:
        return "Head2_NM"
    elif headno==3 :
        return "Head3_NM"
    elif headno==4:
        return "head4_NM"
    elif headno==5:
        return "Head5_NM"
    elif headno==6:
        return "Head6_NM"
    elif headno==7:
        return "Head7_NM"



def writeup(df,name,value):
    try:
        value=value.replace(" USD","")
        if df.shape[0]==3:
            ml=["The ",name," Market was valued at $",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach $",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
            df1=df.drop(index="Total")
            lis=list(df1.index)
            a=df1[df1.columns[0]].idxmax()
            ml+=[a," segment is expected to be the highest contributor to this market, with $",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach $",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
            b=df1[df1.columns[-1]].idxmax()
            ml+=[b," segment is anticipated to reach $",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
            return("".join(ml))
        else:
            ml=["The ",name," Market was valued at $",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach $",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
            df1=df.drop(index="Total")
            lis=list(df1.index)
            a=df1[df1.columns[0]].idxmax()
            ml+=[a," segment is expected to be the highest contributor to this market, with $",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach $",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
            b=df1[df1.columns[-1]].idxmax()
            ml+=[b," segment is anticipated to reach $",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
            df2=df1.drop(index=[a,b])
            c=df2[df2.columns[0]].idxmax()
            d=df2[df2.columns[-1]].idxmax()
            e=round(((df[df.columns[0]][a]+df[df.columns[0]][c])/df[df.columns[0]]["Total"])*100,1)
            ml+=[a," and ",c," segments collectively expected to account for about ",str(e),"% share of the ",name," market in ",df.columns[0],", with the former constituting around ",str(round((df[df.columns[0]][a]/df[df.columns[0]]["Total"])*100,1)),"% share. ",b," and ",d," segments are expected to witness significant growth rates at a CAGR of ",str(round(df[df.columns[-1]][b],2)),"% and ",str(round(df[df.columns[-1]][d],2)),"% respectively, during the forecast period. Presently, share of these two segments is estimated to be around ",str(round(((df[df.columns[0]][b]+df[df.columns[0]][d])/df[df.columns[0]]["Total"])*100,2)),"% in the overall ",name," market in ",df.columns[0],", and is anticipated to reach ",str(round(((df[df.columns[-2]][b]+df[df.columns[-2]][d])/df[df.columns[-2]]["Total"])*100,2)),"% by ",df.columns[-2],". "]
            return("".join(ml))
    except:
        print(df)
        

def writeupvol(df,name,value):
    try:
    
        if df.shape[0]==3:
            ml=["The total demand for ",name," market was ",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach ",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
            df1=df.drop(index="Total")
            lis=list(df1.index)
            a=df1[df1.columns[0]].idxmax()
            ml+=[a," segment is expected to be the highest contributor to this market, with ",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach ",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
            b=df1[df1.columns[-1]].idxmax()
            ml+=[b," segment is anticipated to reach ",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
            return("".join(ml))
        else:
            ml=["The total demand for ",name," market was ",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach ",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
            df1=df.drop(index="Total")
            lis=list(df1.index)
            a=df1[df1.columns[0]].idxmax()
            ml+=[a," segment is expected to be the highest contributor to this market, with ",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach ",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[0]," to ",df.columns[-2],". "]
            b=df1[df1.columns[-1]].idxmax()
            ml+=[b," segment is anticipated to reach ",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
            df2=df1.drop(index=[a,b])
            c=df2[df2.columns[0]].idxmax()
            d=df2[df2.columns[-1]].idxmax()
            e=round(((df[df.columns[0]][a]+df[df.columns[0]][c])/df[df.columns[0]]["Total"])*100,1)
            ml+=[a," and ",c," segments collectively expected to account for about ",str(e),"% share of the ",name," market in ",df.columns[0],", with the former constituting around ",str(round((df[df.columns[0]][a]/df[df.columns[0]]["Total"])*100,1)),"% share. ",b," and ",d," segments are expected to witness significant growth rates at a CAGR of ",str(round(df[df.columns[-1]][b],2)),"% and ",str(round(df[df.columns[-1]][d],2)),"% respectively, during the forecast period. Presently, share of these two segments is estimated to be around ",str(round(((df[df.columns[0]][b]+df[df.columns[0]][d])/df[df.columns[0]]["Total"])*100,2)),"% in the overall ",name," market in ",df.columns[0],", and is anticipated to reach ",str(round(((df[df.columns[-2]][b]+df[df.columns[-2]][d])/df[df.columns[-2]]["Total"])*100,2)),"% by ",df.columns[-2],". "]
            return("".join(ml))
    except:
        print(df)




def reportmake(pdlist,value,volume,market,tablesplit,g):
    dir=os.getcwd()
    doc = docx.Document("Template NM _New.docx")
    # reportname=reportname.split("(1)")[0]
    paragraph_format1 = doc.styles["Table1_Header"].paragraph_format
    paragraph_format2 = doc.styles['Table1_Left'].paragraph_format
    paragraph_format3 = doc.styles['Table1_Right'].paragraph_format
    paragraph_format4 = doc.styles['Table_No_NM'].paragraph_format
    paragraph_format5 = doc.styles['FIG_no_NM'].paragraph_format
    paragraph_format6 = doc.styles['Main_Heading-NM'].paragraph_format
    paragraph_format7 = doc.styles['Head2_NM'].paragraph_format
    paragraph_format8 = doc.styles['Head3_NM'].paragraph_format
    paragraph_format9 = doc.styles['head4_NM'].paragraph_format
    paragraph_format10 = doc.styles['Head5_NM'].paragraph_format
    paragraph_format11 = doc.styles['Head6_NM'].paragraph_format
    paragraph_format12= doc.styles['Head7_NM'].paragraph_format
#     paragraph_format13 = doc.styles['Table_No_NM'].paragraph_format
#     paragraph_format14 = doc.styles['FIG_no_NM'].paragraph_format
    paragraph_format1.keep_with_next = True
    paragraph_format2.keep_with_next = True
    paragraph_format3.keep_with_next = True
    paragraph_format4.keep_with_next = True
    paragraph_format5.keep_with_next = True
    paragraph_format6.keep_with_next = True
    paragraph_format7.keep_with_next = True
    paragraph_format8.keep_with_next = True
    paragraph_format9.keep_with_next = True
    paragraph_format10.keep_with_next = True
    paragraph_format11.keep_with_next = True
    paragraph_format12.keep_with_next = True
    
    dt = datetime.datetime.now()
    t = str(datetime.datetime.timestamp(dt))
    
    Tabno = NamedStyle(name="Tab_no")
    Tabno.font = Font(name=' Franklin Gothic Medium Cond',
                    size=12,
                    bold=True,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')

    heading1 = NamedStyle(name="heading1")
    heading1.font = Font(name='Franklin Gothic Demi Cond',
                    size=19,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')


    heading2= NamedStyle(name="heading2")
    heading2.font = Font(name='Franklin Gothic Medium Cond',
                    size=14,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')


    source=NamedStyle(name="source")
    source.font=Font(name='Franklin Gothic Book',
                    size=7.5,
                    bold=False,
                    italic=True,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='545454')

    main_heading = NamedStyle(name="main_heading")
    main_heading.font = Font(name='Franklin Gothic Medium Cond',
                    size=35,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')
    number=NamedStyle(name="number",number_format="0.00")


    wb = Workbook(write_only=True)
    ws = wb.create_sheet()
    ws.sheet_view.showGridLines = False

    cellmain = WriteOnlyCell(ws)
    cellmain.style = main_heading
    cellmain.alignment = Alignment(horizontal='center', vertical='center')  

    cell = WriteOnlyCell(ws)
    cell.style = Tabno

    cellh1 = WriteOnlyCell(ws)
    cellh1.style = heading1
        

    cellh2 = WriteOnlyCell(ws)
    cellh2.style = heading2

        
    cell1 = WriteOnlyCell(ws)
    cell1.style = source

    cellnum= WriteOnlyCell(ws)
    cellnum.style = number


    def format_row(row, cell):
        for c in row:
            cell.value = c
            yield cell

    styletable = TableStyleInfo(name="TableStyleMedium2", showFirstColumn=True,
                       showLastColumn=True, showRowStripes=True, showColumnStripes=False)

    rnum =1      
    count=1
    count1=1
    count2=1
    tabname=1
    lstcolmain=chr(ord('c')+(g[1]-g[0]+2))

    mainr=[None,None,"Global "+market+" Market"]
    mainrow=format_row(mainr, cellmain)
    ws.append(mainrow)
    ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcolmain+str(rnum))
    rnum+=1


    col=[str(i) for i in range(g[0],g[1]+1)]



    for k in pdlist:
        try:
            f=pd.DataFrame.from_dict(pdlist[k]["data"],orient='index')
            f = f.reindex(columns=col)
            fvol=pd.DataFrame.from_dict(pdlist[k]["vdata"],orient='index')
            fvol =fvol.reindex(columns=col)
            
            mt=(pdlist[k]["val"],", ",f.columns[0],"-",f.columns[-1],", ",value)
            mtvol=(pdlist[k]["vol"],", ",fvol.columns[0],"-",fvol.columns[-1],", ",volume)
            # print(f)
            # print(fvol)
            
            if pdlist[k]["heading"]==1:
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))
                doc.add_paragraph("Overview",style=parastyle(2))
            elif pdlist[k]["heading"]==3 and pdlist[k]["before"]!="":
                doc.add_paragraph(pdlist[k]["before"],style=parastyle(2))
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))
            elif pdlist[k]["heading"]==5 and pdlist[k]["before"]!="":
                doc.add_paragraph(pdlist[k]["before"],style=parastyle(4))
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))
            else:
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))



            # p=doc.add_paragraph("".join(mt),style="FIG_no_NM")
        
            f.loc["Total"] = f.sum()
            f['CAGR % '+"("+f.columns[1]+"-"+f.columns[-1]+")" ] = ((f.iloc[:, -1].div(f.iloc[:, 1]).pow(1./(len(f.columns) - 2)).sub(1))*100 )    

            fvol.loc["Total"] = fvol.sum()
            fvol['CAGR % '+"("+fvol.columns[1]+"-"+fvol.columns[-1]+")" ] = ((fvol.iloc[:, -1].div(fvol.iloc[:, 1]).pow(1./(len(fvol.columns) - 2)).sub(1))*100 )   
            if volume=="Units":
                fvol = fvol.round().astype({i:'int' for i in col })
                # print(fvol)

            lstcol=chr(ord('c')+f.shape[1])
            if pdlist[k]["heading"]==1:
                h1=[None,None,str(count1)+". "+(k.split("#")[0]).upper()]
                count1+=1
                count2=1
                head_row = format_row(h1, cellh1)
                ws.append(head_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1  
                        
                
            elif pdlist[k]["heading"]==3 and pdlist[k]["before"]!="":
                h2=[None,None,str(count2)+". "+(pdlist[k]["before"]).upper()]
                count2+=1
                head2_row = format_row(h2, cellh2)
                ws.append(head2_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1  

                
                

            lstcol=chr(ord('c')+f.shape[1])
            r1=[None,None,"TABLE "+str(count)+". "+("".join(mt)).upper()]
            count+=1
            row1=[None,None] 
            first_row = format_row(r1, cell)
            ws.append(first_row)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            rnum +=1  
            rowtable=rnum
            df=list(dataframe_to_rows(f, index=True, header=True))
            for i in range (len(df)):
                row1=[None,None]
                row= row1+df[i]
            #         print((row))
                if i==0:
                    row=list(map(str, df[i]))
                    row[0]=pdlist[k]["comp"]
                    row=row1+row
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  
                elif i==1:
                    continue
                else:
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  


            tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
            tabname+=1
            headings = [pdlist[k]["comp"]]+list(f.columns)
            tab1._initialise_columns()
            for column, value1 in zip(tab1.tableColumns, headings):
                column.name = value1
        #     tab1.headerRowCount = False
            
            tab1.tableStyleInfo = styletable
            ws.add_table(tab1)
            sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
            sourcerow = format_row(sourcer, cell1)
            ws.append(sourcerow)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            ws.append([None])
            ws.append([None])
            rnum +=3

            r1=[None,None,"TABLE "+str(count)+". "+("".join(mtvol)).upper()]
            count+=1
            row1=[None,None] 
            first_row = format_row(r1, cell)
            ws.append(first_row)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            rnum +=1  
            rowtable=rnum
            df=list(dataframe_to_rows(fvol, index=True, header=True))
            for i in range (len(df)):
                row1=[None,None]
                row= row1+df[i]
            #         print((row))
                if i==0:
                    row=list(map(str, df[i]))
                    row[0]=pdlist[k]["comp"]
                    row=row1+row
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  
                elif i==1:
                    continue
                else:
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  
            
            tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
            tabname+=1
            headings = [pdlist[k]["comp"]]+list(f.columns)
            tab1._initialise_columns()
            for column, value1 in zip(tab1.tableColumns, headings):
                column.name = value1
        #     tab1.headerRowCount = False
            
            tab1.tableStyleInfo = styletable
            ws.add_table(tab1)

            sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
            sourcerow = format_row(sourcer, cell1)
            ws.append(sourcerow)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            ws.append([None])
            ws.append([None])
            rnum +=3



            
        

            if (tablesplit):
                f=f.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)
                fvol=fvol.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)

            p=doc.add_paragraph("".join(mt),style="FIG_no_NM")
            f1=f.drop([f.columns[-1]],axis=1)
            f1=f1.drop(index=("Total"))
            pic=plot(f1.T,t)
            doc.add_picture(pic,width=Cm(16.98), height=Cm(8.51))
            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")
            # doc.add_paragraph(" ")



            p=doc.add_paragraph("".join(mt),style="Table_No_NM")
            table = doc.add_table(rows=f.shape[0]+1, cols=f.shape[1]+1)
            table.style = 'Table_Style_NM'
            table_cells = table._cells
            for r in range (len(table.rows)):
                row=table.rows[r]
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                if r==0:
                    row.height = Cm(1.1)
                else:
                    row.height = Cm(0.7)

            table.cell(0,0).paragraphs[0].style="Table1_Header"
            table.cell(0,0).width = 1242000.2592
            table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
            table.columns[0].width= Cm(3.45)
            table.columns[-1].width=Cm(1.97)


            for k1 in range(1,len(f.columns)+1):
                table.cell(0,k1).paragraphs[0].style="Table1_Header"
                table.cell(0,k1).paragraphs[0].text=(f.columns[k1-1])
            for i in range(1,f.shape[0]+1):
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].style="Table1_Left"
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].text=(f.index[i-1])
                for j in range(1,f.shape[1]+1):
                    table_cells[j + i * (f.shape[1]+1)].paragraphs[0].style="Table1_Right"
                    if (j %(f.shape[1]))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2))+"%")
                        
                    elif (j %(f.shape[1]+1))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 1242000.2592
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))
                        
                    else:
                        table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))

            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis",style="Source_NM")
            


            doc.add_paragraph(writeup(f,pdlist[k]["name"],value),style="Normal")
            
            
            
            p=doc.add_paragraph("".join(mtvol),style="FIG_no_NM")
            f1vol=fvol.drop([fvol.columns[-1]],axis=1)
            f1vol=f1vol.drop(index=("Total"))
            picv=plot(f1vol.T,t)
            doc.add_picture(picv,width=Cm(16.98), height=Cm(8.51))
            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")
            # doc.add_paragraph(" ")

            
            p=doc.add_paragraph("".join(mtvol),style="Table_No_NM")
            table = doc.add_table(rows=fvol.shape[0]+1, cols=fvol.shape[1]+1)
            table.style = 'Table_Style_NM'
            table_cells = table._cells
            for r in range (len(table.rows)):
                row=table.rows[r]
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                if r==0:
                    row.height = Cm(1.1)
                else:
                    row.height = Cm(0.7)

            table.cell(0,0).paragraphs[0].style="Table1_Header"
            table.cell(0,0).width = 1242000.2592
            table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
            table.columns[0].width= Cm(3.45)
            table.columns[-1].width=Cm(1.97)

            for k1 in range(1,len(f.columns)+1):
                table.cell(0,k1).paragraphs[0].style="Table1_Header"
                table.cell(0,k1).paragraphs[0].text=(fvol.columns[k1-1])
            for i in range(1,f.shape[0]+1):
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].style="Table1_Left"
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].text=(fvol.index[i-1])
                for j in range(1,f.shape[1]+1):
                    table_cells[j + i * (f.shape[1]+1)].paragraphs[0].style="Table1_Right"
                    if (j %(f.shape[1]))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2))+"%")
                        
                    elif (j %(f.shape[1]+1))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 1242000.2592
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2)))
                        
                    else:
                        table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(fvol.values[i-1][j-1],2)))

            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis",style="Source_NM")
            
            
            doc.add_paragraph(writeupvol(fvol,pdlist[k]["name"],volume),style="Normal")
        except Exception as e:
            dir= os.getcwd()
            logfile=dir+"\error_log.txt"
            file=open(logfile,"a")
            
            # data1={data}
            file.write("/n"+market.upper()+" "+str(e))
            print(e)
            # f= open(".txt","a")

            continue


    doc.save(dir+"/files/"+"Global "+market+" Market Value and Volume Report Forecast till "+str(g[1])+".docx")
    wb.save(dir+"/files/"+"Global "+market+" Market Value and Volume Excel Report Forecast till "+str(g[1])+".xlsx")
    os.remove(pic)  

    return(("/files/"+"Global "+market+" Market Value and Volume Report Forecast till "+str(g[1])+".docx","/files/"+"Global "+market+" Market Value and Volume Excel Report Forecast till "+str(g[1])+".xlsx"))





def det(toc):
    reg={}
    tocl1={}
    for i in toc:
        count=[]
        for j in i["tocreg"]:
            count.append(j["country"])
        reg[i["region"]]=count
    for i in toc[0]["tocreg"][0]["toc"]:
        tocl2={}
        for j in i["toc_component"]:
            tocl3={}
            if j["toc2_component"]!=[]:
                for k in j["toc2_component"]:
                    tocl4=[]
                    if k["toc3_component"]!=[]:
                        for k1 in k["toc3_component"]:
                            tocl4.append(k1["toc4_name"])
                    tocl3[k["toc3_name"]]=tocl4
            tocl2[j["toc2_name"]]=tocl3
        tocl1[i["toc_name"]]=tocl2

    return((reg,tocl1))


# import json

def final(toc1,year,ASP,market,value,volume,tablesplit):
    market=market.split("(1)")[0]
    v=list(year.keys())
    g=[int(v[0]),int(v[-1])]
    for i in year:
        sk=year[i]
        year[i]=sk.replace(" ","")
        year[i]=sk.replace(",","")
    for n in ASP:
        sk=ASP[n]
        ASP[n]=sk.replace(" ","")
        ASP[n]=sk.replace(",","")
    
    if volume=="Units":
        yearvol={i:int(round(float(ASP[i]))) for i in ASP}
    else:
        yearvol=ASP
    print(yearvol)
    global100(toc1)
    region100(toc1)
    toc100(toc1)
    global1(toc1,year,yearvol,g)
    
    f=det(toc1)
    j=[]
    for i in f[0]:
        j+=(f[0][i])
    regdet={"region":list(f[0].keys()),"country":j}

    pdlist=makepdlist(toc1,f[1],f[0],market)
    contchange=f[0]
    for i in contchange :
        contchange[i]=[{"label":j,"country":j,"value":j} for j in contchange[i]]
    return(reportmake(pdlist,value,volume,market,tablesplit,g),regdet,contchange)
    