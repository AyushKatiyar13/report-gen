import pandas as pd
import copy
from collections import Counter
import matplotlib.pyplot as plt
import docx
from docx.shared import Inches,Cm
import datetime
import os
from docx.enum.table import WD_ROW_HEIGHT_RULE
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, Border, Side, NamedStyle,Alignment
from openpyxl.cell.cell import WriteOnlyCell
from openpyxl.worksheet.table import Table, TableStyleInfo
from mpmath import *                            


def auto100(regen):
    regend=[]
    regen=[mpf(str(i).replace("%","")) for i in regen]
    sm=sum(regen)
    
    if sm!=100.0:
        if sm>100.0:
            diff=mpf((sm-100.0)/len(regen))
            for i in regen:
                regend.append(i-diff)
        else :
            diff=mpf((100.0-sm)/len(regen))
            for i in regen:
                regend.append(i+diff)
    else:
        regend=regen
    return(regend)

       
# def region100(toc):
#     # for i in toc:
#     spstart=[]
#     spend=[]
#     for i in toc:
#         start=mpf((str(i["split"][0])).replace("%",""))
#         end=mpf((str(i["split"][1])).replace("%",""))
#         spstart.append(start)
#         spend.append(end)
        
#     regst1=auto100(spstart)
#     regen1=auto100(spend)
#     for i in range(len(toc)):
#         toc[i]["split"][0]=regst1[i]
#         toc[i]["split"][1]=regen1[i]
#         # i["tocreg"][j]["split"][0]=regst1[j]
#         # i["tocreg"][j]["split"][1]=regen1[j]


       
def region100(toc):
    # for i in toc:
    spstart=[]
    spend=[]
    spstartvol=[]
    spendvol=[]
    for i in toc:
        start=mpf((str(i["split"][0])).replace("%",""))
        end=mpf((str(i["split"][1])).replace("%",""))
        startvol=mpf((str(i["vsplit"][0])).replace("%",""))
        endvol=mpf((str(i["vsplit"][1])).replace("%",""))
        spstart.append(start)
        spend.append(end)
        spstartvol.append(startvol)
        spendvol.append(endvol)
        
    regst1=auto100(spstart)
    regen1=auto100(spend)
    regst1vol=auto100(spstartvol)
    regen1vol=auto100(spendvol)
    for i in range(len(toc)):
        toc[i]["split"][0]=regst1[i]
        toc[i]["split"][1]=regen1[i]
        toc[i]["vsplit"][0]=regst1vol[i]
        toc[i]["vsplit"][1]=regen1vol[i]



def toc100(toc):
    for i in toc:
        for k in i["toc"]: 
            sttc2=[]
            endtc2=[]
            sttc2vol=[]
            endtc2vol=[]
            for tc2 in k["toc_component"]:
                sttc2.append(tc2["split"][0])
                endtc2.append(tc2["split"][1])
                sttc2vol.append(tc2["vsplit"][0])
                endtc2vol.append(tc2["vsplit"][1])
                if tc2["toc2_component"]!=[]:
                    sttc3=[]
                    endtc3=[]
                    sttc3vol=[]
                    endtc3vol=[]
                    for tc3 in tc2["toc2_component"]:
                        sttc3.append(tc3["split"][0])
                        endtc3.append(tc3["split"][1])
                        sttc3vol.append(tc3["vsplit"][0])
                        endtc3vol.append(tc3["vsplit"][1])
                        if tc3["toc3_component"]!=[]:
                            sttc4=[]
                            endtc4=[]
                            sttc4vol=[]
                            endtc4vol=[]
                            for tc4 in tc3["toc3_component"]:
                                # print(tc4)
                                sttc4.append(tc4["split"][0])
                                endtc4.append(tc4["split"][1])
                                sttc4vol.append(tc4["vsplit"][0])
                                endtc4vol.append(tc4["vsplit"][1])
                            
                            
                            regst13=auto100(sttc4)
                            regen13=auto100(endtc4)
                            regst13vol=auto100(sttc4vol)
                            regen13vol=auto100(endtc4vol)
                            for tc4 in range(len(tc3["toc3_component"])):
                                tc3["toc3_component"][tc4]["split"][0]=regst13[tc4]
                                tc3["toc3_component"][tc4]["split"][1]=regen13[tc4]
                                tc3["toc3_component"][tc4]["vsplit"][0]=regst13vol[tc4]
                                tc3["toc3_component"][tc4]["vsplit"][1]=regen13vol[tc4]
                            

                    regst12=auto100(sttc3)
                    regen12=auto100(endtc3)
                    regst12vol=auto100(sttc3vol)
                    regen12vol=auto100(endtc3vol)
                    for tc3 in range(len(tc2["toc2_component"])):
                        tc2["toc2_component"][tc3]["split"][0]=regst12[tc3]
                        tc2["toc2_component"][tc3]["split"][1]=regen12[tc3]
                        tc2["toc2_component"][tc3]["vsplit"][0]=regst12vol[tc3]
                        tc2["toc2_component"][tc3]["vsplit"][1]=regen12vol[tc3]


        regst1=auto100(sttc2)
        regen1=auto100(endtc2)  
        regst1vol=auto100(sttc2vol)
        regen1vol=auto100(endtc2vol)  
        for tc2 in range(len(k["toc_component"])):
            k["toc_component"][tc2]["split"][0]=regst1[tc2]
            k["toc_component"][tc2]["split"][1]=regen1[tc2]
            k["toc_component"][tc2]["vsplit"][0]=regst1vol[tc2]
            k["toc_component"][tc2]["vsplit"][1]=regen1vol[tc2]
            




                    
def region(toc,year,yearvol,g):
    # print(toc)
    for i in toc:
        s1=mpf(str(i["split"][0]).replace("%",""))
        s2=mpf(str(i["split"][1]).replace("%",""))
        s1v=mpf(str(i["vsplit"][0]).replace("%",""))
        s2v=mpf(str(i["vsplit"][1]).replace("%",""))
        step=((s2)-(s1))/(int(g[1])-int(g[0]))
        stepv=(s2v-s1v)/(int(g[1])-int(g[0]))
        i["split"]={}
        i["vsplit"]={}
        for k in range (g[0],g[1]+1):
            i["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
            i["vsplit"][str(k)]=float((s1v*mpf(str(year[str(k)]).replace(",","")))/100)
            s1+=step
            s1v+=stepv
            

def tocl2reg(toc,year1,year1vol,g):
    for i in toc:
        for j in (i["toc_component"]):
            s1=mpf(str(j["split"][0]).replace("%",""))
            s2=mpf(str(j["split"][1]).replace("%",""))
            s1v=mpf(str(j["vsplit"][0]).replace("%",""))
            s2v=mpf(str(j["vsplit"][1]).replace("%",""))
            step=((s2)-(s1))/(int(g[1])-int(g[0]))
            stepv=((s2v)-(s1v))/(int(g[1])-int(g[0]))
            j["split"]={}
            j["vsplit"]={}
            for k in range (g[0],g[1]+1):
                j["split"][str(k)]=float((s1*mpf(str(year1[str(k)]).replace(",","")))/100)
                j["vsplit"][str(k)]=float((s1v*mpf(str(year1vol[str(k)]).replace(",","")))/100)
                # float((s1v*mpf(year1vol[str(k)]))/100)
                s1+=step
                s1v+=stepv
    
def tocl3reg(toc,g):
    for j in range(len(toc)):
        for c in toc[j]["toc_component"]:
            year=c["split"]
            yearvol=c["vsplit"]
            if (len(c["toc2_component"]))!=0:
                for t in (c["toc2_component"]):
                    s1=mpf(str(t["split"][0]).replace("%",""))
                    s2=mpf(str(t["split"][1]).replace("%",""))
                    s1v=mpf(str(t["vsplit"][0]).replace("%",""))
                    s2v=mpf(str(t["vsplit"][1]).replace("%",""))
                    step=((s2)-(s1))/(int(g[1])-int(g[0]))
                    stepv=((s2v)-(s1v))/(int(g[1])-int(g[0]))
                    t["split"]={}
                    t["vsplit"]={}
                    for k in range (g[0],(g[1]+1)):
                        t["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
                        t["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
                        # loat((s1v*mpf(yearvol[str(k)]))/100)
                        s1+=step
                        s1v+=stepv
            
def tocl4reg(toc,g):
    for j in range(len(toc)):
        for c in toc[j]["toc_component"]:
            for t in (c["toc2_component"]):
                year=t["split"]
                yearvol=t["vsplit"]
                if (len(t["toc3_component"]))!=0:
                    for r in (t["toc3_component"]):
                        s1=mpf(str(r["split"][0]).replace("%",""))
                        s2=mpf(str(r["split"][1]).replace("%",""))
                        s1v=mpf(str(r["vsplit"][0]).replace("%",""))
                        s2v=mpf(str(r["vsplit"][1]).replace("%",""))
                        step=((s2)-(s1))/(int(g[1])-int(g[0]))
                        r["split"]={}
                        stepv=((s2v)-(s1v))/(int(g[1])-int(g[0]))
                        r["vsplit"]={}
                        for k in range (g[0],(g[1]+1)):
                            r["split"][str(k)]=float((s1*mpf(str(year[str(k)]).replace(",","")))/100)
                            # float((s1*mpf(year[str(k)]))/100)
                            s1+=step
                            r["vsplit"][str(k)]=float((s1v*mpf(str(yearvol[str(k)]).replace(",","")))/100)
                            # float((s1v*mpf(yearvol[str(k)]))/100)
                            s1v+=stepv
                                                                
def tocchanger(toc,year,yearvol,g):
    region(toc,year,yearvol,g)
    for i in toc:
        year1=i["split"]
        yearvol1=i["vsplit"]
        tocl2reg(i["toc"],year1,yearvol1,g)
        tocl3reg(i["toc"],g)
        tocl4reg(i["toc"],g)
    return toc


def det(toc):
    reg=[]
    tocl1={}
    for i in toc:
        reg.append(i["country"])

    for i in toc[0]["toc"]:
        tocl2={}
        for j in i["toc_component"]:
            tocl3={}
            if j["toc2_component"]!=[]:
                for k in j["toc2_component"]:
                    tocl4=[]
                    if k["toc3_component"]!=[]:
                        for k1 in k["toc3_component"]:
                            tocl4.append(k1["toc4_name"])
                    tocl3[k["toc3_name"]]=tocl4
            tocl2[j["toc2_name"]]=tocl3
        tocl1[i["toc_name"]]=tocl2
    return((reg,tocl1))



def regionbytocl2(toc,l1,l2): 
    k1=Counter()
    for i in toc:
        for j in i["toc"]:
            if j["toc_name"]==l1:
                for n in j["toc_component"]:
                    if  n["toc2_name"]==l2:
                        k1+=n["split"]            
    return(k1)

def regionbytocl3(toc,l1,l2,l3):
    k1=Counter()
    for i in toc:
        for j in i["toc"]:
            if j["toc_name"]==l1:
                for n in j["toc_component"]:
                    if  n["toc2_name"]==l2:
                        for n1 in n["toc2_component"]:
                            if n1["toc3_name"]==l3:
                                k1+=n1["split"]
    return(k1)

def regionbytocl4(toc,l1,l2,l3,l4):
    k1=Counter()
    for i in toc:
        for j in i["toc"]:
            if j["toc_name"]==l1:
                for n in j["toc_component"]:
                    if  n["toc2_name"]==l2:
                        for n1 in n["toc2_component"]:
                            if n1["toc3_name"]==l3:
                                for n2 in n1["toc3_component"]:
                                    if n2["toc4_name"]==l4:
                                        k1+=n2["split"]
    return(k1)

def regionbytocl2vol(toc,l1,l2):
    k1=Counter()
    for i in toc:
        for j in i["toc"]:
            if j["toc_name"]==l1:
                for n in j["toc_component"]:
                    if  n["toc2_name"]==l2:
                        k1+=n["vsplit"]            
    return(k1)

def regionbytocl3vol(toc,l1,l2,l3):
    k1=Counter()
    for i in toc:
        for j in i["toc"]:
            if j["toc_name"]==l1:
                for n in j["toc_component"]:
                    if  n["toc2_name"]==l2:
                        for n1 in n["toc2_component"]:
                            if n1["toc3_name"]==l3:
                                k1+=n1["vsplit"]
    return(k1)

def regionbytocl4vol(toc,l1,l2,l3,l4):
    k1=Counter()
    for i in toc:
        for j in i["toc"]:
            if j["toc_name"]==l1:
                for n in j["toc_component"]:
                    if  n["toc2_name"]==l2:
                        for n1 in n["toc2_component"]:
                            if n1["toc3_name"]==l3:
                                for n2 in n1["toc3_component"]:
                                    if n2["toc4_name"]==l4:
                                        k1+=n2["vsplit"]
    return(k1)


def countrybytocl2vol(toc,countr,l1,l2):
    for i in toc:
        if i["country"]==countr:
            for j in i["toc"]:
                if j["toc_name"]==l1:
                    for n in j["toc_component"]:
                        if  n["toc2_name"]==l2:
                            return(n["vsplit"])          

def countrybytocl3vol(toc,countr,l1,l2,l3):
    for i in toc:
        if i["country"]==countr:
            for j in i["toc"]:
                if j["toc_name"]==l1:
                    for n in j["toc_component"]:
                        if  n["toc2_name"]==l2:
                            for n1 in n["toc2_component"]:
                                if n1["toc3_name"]==l3:
                                    return(n1["vsplit"])

def countrybytocl4vol(toc,countr,l1,l2,l3,l4):
    for i in toc:
        if i["country"]==countr:
            for j in i["toc"]:
                if j["toc_name"]==l1:
                    for n in j["toc_component"]:
                        if  n["toc2_name"]==l2:
                            for n1 in n["toc2_component"]:
                                if n1["toc3_name"]==l3:
                                    for n2 in n1["toc3_component"]:
                                        if n2["toc4_name"]==l4:
                                            return(n2["vsplit"])

def countrybytocl2(toc,countr,l1,l2):
    for i in toc:
        if i["country"]==countr:
            for j in i["toc"]:
                if j["toc_name"]==l1:
                    for n in j["toc_component"]:
                        if  n["toc2_name"]==l2:
                            return(n["split"])          

def countrybytocl3(toc,countr,l1,l2,l3):
    for i in toc:
        if i["country"]==countr:
            for j in i["toc"]:
                if j["toc_name"]==l1:
                    for n in j["toc_component"]:
                        if  n["toc2_name"]==l2:
                            for n1 in n["toc2_component"]:
                                if n1["toc3_name"]==l3:
                                    return(n1["split"])

def countrybytocl4(toc,countr,l1,l2,l3,l4):
    for i in toc:
        if i["country"]==countr:
            for j in i["toc"]:
                if j["toc_name"]==l1:
                    for n in j["toc_component"]:
                        if  n["toc2_name"]==l2:
                            for n1 in n["toc2_component"]:
                                if n1["toc3_name"]==l3:
                                    for n2 in n1["toc3_component"]:
                                        if n2["toc4_name"]==l4:
                                            return(n2["split"])
                                        
def marketbycountry(toc,countr):
    for i in toc:
        if i["country"]==countr:
            return i["split"]


                                       
def marketbycountryvol(toc,countr):
    for i in toc:
        if i["country"]==countr:
            return i["vsplit"]


def pdlistmaker(toc,tocl1,reg,market,reportname):
    pdlist={}
    a=0
    b=0
    c=0
    count=1
    count1=0
    count2=0
    for l1 in tocl1:
        pdlist[market+" "+reportname+" MARKET, BY "+l1+"#"+str(count)]={"heading":1,"comp":l1,"wrtup":market+" "+reportname,"data":{},"vdata":{},"val":market+" "+reportname+" VALUE MARKET, BY "+l1,"vol":market+" "+reportname+" VOLUME MARKET, BY "+l1}
        for l2 in tocl1[l1]:
            pdlist[market+" "+reportname+" MARKET, BY "+l1+"#"+str(count)]["data"][l2]=regionbytocl2(toc,l1,l2)
            pdlist[market+" "+reportname+" MARKET, BY "+l1+"#"+str(count)]["vdata"][l2]=regionbytocl2vol(toc,l1,l2)
            if tocl1[l1][l2]!={}:
                pdlist[market+" "+l2+" MARKET, BY "+l1+"#"+str(count1)]={"heading":3,"before":l2,"comp":l2,"wrtup":market+" "+l2,"data":{},"vdata":{},"val":market+" "+l2+" VALUE MARKET, BY "+l1,"vol":market+" "+l2+" VOLUME MARKET, BY "+l1}
                a=1
                for l3 in tocl1[l1][l2]:
                    pdlist[market+" "+l2+" MARKET, BY "+l1+"#"+str(count1)]["data"][l3]=regionbytocl3(toc,l1,l2,l3)
                    pdlist[market+" "+l2+" MARKET, BY "+l1+"#"+str(count1)]["vdata"][l3]=regionbytocl3vol(toc,l1,l2,l3)
                    count2=count1
                    if tocl1[l1][l2][l3]!=[]:
                        pdlist[market+" "+l3+" MARKET, BY "+l2+"#"+str(count2)]={"heading":5,"before":l3,"comp":l3,"wrtup":market+" "+l3,"data":{},"vdata":{},"val":market+" "+l3+" VALUE MARKET, BY "+l2,"vol":market+" "+l3+" VOLUME MARKET, BY "+l2}
                        b=1
                        for l4 in tocl1[l1][l2][l3]:
                            pdlist[market+" "+l3+" MARKET, BY "+l2+"#"+str(count2)]["data"][l4]=regionbytocl4(toc,l1,l2,l3,l4)
                            pdlist[market+" "+l3+" MARKET, BY "+l2+"#"+str(count2)]["vdata"][l4]=regionbytocl4vol(toc,l1,l2,l3,l4)
                            pdlist[market+" "+l4+" MARKET, BY COUNTRY"+"#"+str(count2)]={"heading":6,"before":l4,"comp":"Country","wrtup":market+" "+l4,"data":{},"vdata":{},"val":market+" "+l4+" VALUE MARKET, BY COUNTRY","vol":market+" "+l4+" VOLUME MARKET, BY COUNTRY"}
                            c=1
                            for countr in reg:
                                pdlist[market+" "+l4+" MARKET, BY COUNTRY"+"#"+str(count2)]["data"][countr]=countrybytocl4(toc,countr,l1,l2,l3,l4)
                                pdlist[market+" "+l4+" MARKET, BY COUNTRY"+"#"+str(count2)]["vdata"][countr]=countrybytocl4vol(toc,countr,l1,l2,l3,l4)
                        count2+=1
                    if b==0:
                        pdlist[market+" "+l3+" MARKET, BY COUNTRY"+"#"+str(count1)]={"heading":4,"before":l3,"comp":"Country","wrtup":market+" "+l3,"data":{},"vdata":{},"val":market+" "+l3+" VALUE MARKET, BY COUNTRY","vol":market+" "+l3+" VOLUME MARKET, BY COUNTRY"}
                    else:
                        pdlist[market+" "+l3+" MARKET, BY COUNTRY"+"#"+str(count1)]={"heading":4,"before":"","comp":"Country","wrtup":market+" "+l3,"data":{},"vdata":{},"val":market+" "+l3+" VALUE MARKET, BY COUNTRY","vol":market+" "+l3+" VOLUME MARKET, BY COUNTRY"}
                    for countr in reg:
                        pdlist[market+" "+l3+" MARKET, BY COUNTRY"+"#"+str(count1)]["data"][countr]=countrybytocl3(toc,countr,l1,l2,l3)
                        pdlist[market+" "+l3+" MARKET, BY COUNTRY"+"#"+str(count1)]["vdata"][countr]=countrybytocl3vol(toc,countr,l1,l2,l3)
            if a==0:
                pdlist[market+" "+l2+" MARKET, BY COUNTRY"+"#"+str(count)]={"heading":3,"before":l2,"comp":"Country","wrtup":market+" "+l2,"data":{},"vdata":{},"val":market+" "+l2+" VALUE MARKET, BY COUNTRY","vol":market+" "+l2+" VOLUME MARKET, BY COUNTRY"}
            else:
                pdlist[market+" "+l2+" MARKET, BY COUNTRY"+"#"+str(count)]={"heading":3,"before":"","comp":"Country","wrtup":market+" "+l2,"data":{},"vdata":{},"val":market+" "+l2+" VALUE MARKET, BY COUNTRY","vol":market+" "+l2+" VOLUME MARKET, BY COUNTRY"}

            for countr in reg:
                pdlist[market+" "+l2+" MARKET, BY COUNTRY"+"#"+str(count)]["data"][countr]=countrybytocl2(toc,countr,l1,l2)
                pdlist[market+" "+l2+" MARKET, BY COUNTRY"+"#"+str(count)]["vdata"][countr]=countrybytocl2vol(toc,countr,l1,l2)
        count=count1+count+count2



    pdlist[market+" "+reportname+" MARKET, BY COUNTRY"+"#"+str(count)]={"heading":1,"comp":"Country","wrtup":market+" "+reportname,"data":{},"vdata":{},"val":market+" "+reportname+" VALUE MARKET, BY COUNTRY","vol":market+" "+reportname+" VOLUME MARKET, BY COUNTRY"}
    for countr in reg:
        pdlist[market+" "+reportname+" MARKET, BY COUNTRY"+"#"+str(count)]["data"][countr]=marketbycountry(toc,countr)
        pdlist[market+" "+reportname+" MARKET, BY COUNTRY"+"#"+str(count)]["vdata"][countr]=marketbycountryvol(toc,countr)
        c=0
        for l1 in tocl1:
            if c==0:
                pdlist[countr+" "+reportname+" MARKET, BY "+l1+"#"+str(count2)]={"heading":3,"before":countr,"comp":l1,"wrtup":countr+" "+reportname,"data":{},"vdata":{},"val":countr+" "+reportname+"MARKET, BY "+l1,"vol":countr+" "+reportname+"MARKET, BY "+l1}
                c=1
            else:
                pdlist[countr+" "+reportname+" MARKET, BY "+l1+"#"+str(count2)]={"heading":3,"before":"","comp":l1,"wrtup":countr+" "+reportname,"data":{},"vdata":{},"val":countr+" "+reportname+"MARKET, BY "+l1,"vol":countr+" "+reportname+"MARKET, BY "+l1}
            for l2 in tocl1[l1]:
                pdlist[countr+" "+reportname+" MARKET, BY "+l1+"#"+str(count2)]["data"][l2]=countrybytocl2(toc,countr,l1,l2)
                pdlist[countr+" "+reportname+" MARKET, BY "+l1+"#"+str(count2)]["vdata"][l2]=countrybytocl2vol(toc,countr,l1,l2)
                if (tocl1[l1][l2])!={}:
                    pdlist[countr+" "+l2+" MARKET, BY "+l1+"#"+str(count1)]={"heading":4,"before":"","comp":l2,"wrtup":countr+" "+l2,"data":{},"vdata":{},"val":countr+" "+l2+" VALUE MARKET, BY "+l1,"vol":countr+" "+l2+" VOLUME MARKET, BY "+l1}
                    for l3 in (tocl1[l1][l2]):
                        pdlist[countr+" "+l2+" MARKET, BY "+l1+"#"+str(count1)]["data"][l3]=countrybytocl3(toc,countr,l1,l2,l3)
                        pdlist[countr+" "+l2+" MARKET, BY "+l1+"#"+str(count1)]["vdata"][l3]=countrybytocl3vol(toc,countr,l1,l2,l3)
                        if (tocl1[l1][l2][l3])!=[]:
                            count21=count2+count1
                            pdlist[countr+" "+l3+" MARKET, BY "+l2+"#"+str(count21)]={"heading":5,"before":"","comp":l3,"wrtup":countr+" "+l3,"data":{},"vdata":{},"val":countr+" "+l3+" VALUE MARKET, BY "+l2,"vol":countr+" "+l3+" VOLUME MARKET, BY "+l2}
                            for l4 in (tocl1[l1][l2][l3]):
                                pdlist[countr+" "+l3+" MARKET, BY "+l2+"#"+str(count21)]["data"][l4]=countrybytocl4(toc,countr,l1,l2,l3,l4)
                                pdlist[countr+" "+l3+" MARKET, BY "+l2+"#"+str(count21)]["vdata"][l4]=countrybytocl4vol(toc,countr,l1,l2,l3,l4)
        count2+=1
                                
    return(pdlist)



def plot(j,t):
    dir=os.getcwd()
    ax=j.plot(kind="bar",figsize=(9,6))
    ax.set_axisbelow(True)
    ax.yaxis.grid(color='gray', linestyle='dashed')
    plt.xticks(rotation='horizontal')
    plt.legend(loc="upper left")
    plt.savefig(dir+"/temp/"+t+'example.png',bbox_inches='tight',pad_inches = 0)
    plt.close()
    return(dir+"/temp/"+t+'example.png')


def parastyle(headno):
    if headno==1:
        return "Main_Heading-NM"
    elif headno==2:
        return "Head2_NM"
    elif headno==3 :
        return "Head3_NM"
    elif headno==4:
        return "head4_NM"
    elif headno==5:
        return "Head5_NM"
    elif headno==6:
        return "Head6_NM"



def writeup(df,name,value):
    value=value.replace(" USD","")
    if df.shape[0]==3:
        ml=["The ",name," Market was valued at $",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach $",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with $",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach $",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        b=df1[df1.columns[-1]].idxmax()
        ml+=[b," segment is anticipated to reach $",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
        return("".join(ml))
    else:
        ml=["The ",name," Market was valued at $",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach $",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with $",str(round(df[df.columns[0]][a],2))," ",value," in ",df.columns[0],", and is anticipated to reach $",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        b=df1[df1.columns[-1]].idxmax()
        ml+=[b," segment is anticipated to reach $",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
        df2=df1.drop(index=[a,b])
        c=df2[df2.columns[0]].idxmax()
        d=df2[df2.columns[-1]].idxmax()
        e=round(((df[df.columns[0]][a]+df[df.columns[0]][c])/df[df.columns[0]]["Total"])*100,1)
        ml+=[a," and ",c," segments collectively expected to account for about ",str(e),"% share of the",name," market in ",df.columns[0],", with the former constituting around ",str(round((df[df.columns[0]][a]/df[df.columns[0]]["Total"])*100,1)),"% share. ",b," and ",d," segments are expected to witness significant growth rates at a CAGR of ",str(round(df[df.columns[-1]][b],2)),"% and ",str(round(df[df.columns[-1]][d],2)),"% respectively, during the forecast period. Presently, share of these two segments is estimated to be around ",str(round(((df[df.columns[0]][b]+df[df.columns[0]][d])/df[df.columns[0]]["Total"])*100,2)),"% in the overall U.S. ",name," market in ",df.columns[0],", and is anticipated to reach $",str(round(((df[df.columns[-2]][b]+df[df.columns[-2]][d])/df[df.columns[-2]]["Total"])*100,2)),"% by ",df.columns[-2],". "]
        return("".join(ml))
        


def writeupvol(df,name,value):
    
    if df.shape[0]==3:
        ml=["The total demand for ",name," market was ",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach ",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with ",str(round(df[df.columns[0]][a],2))," in ",df.columns[0],", and is anticipated to reach ",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        b=df1[df1.columns[-1]].idxmax()
        ml+=[b," segment is anticipated to reach ",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
        return("".join(ml))
    else:
        ml=["The total demand for ",name," market was ",str(round(df[df.columns[0]]["Total"],2))," ",value," in ",df.columns[0],", and is projected to reach ",str(round(df[df.columns[-2]]["Total"],2))," ",value," by " ,df.columns[-2]," growing at a CAGR of ",str(round(df[df.columns[-1]]["Total"],2)),"% from " ,df.columns[1]," to ",df.columns[-2],". "]
        df1=df.drop(index="Total")
        lis=list(df1.index)
        a=df1[df1.columns[0]].idxmax()
        ml+=[a," segment is expected to be the highest contributor to this market, with ",str(round(df[df.columns[0]][a],2))," in ",df.columns[0],", and is anticipated to reach ",str(round(df[df.columns[-2]][a],2))," ",value," by " ,df.columns[-2],", registering a CAGR of ",str(round(df[df.columns[-1]][a],2)),"% from " ,df.columns[0]," to ",df.columns[-2],". "]
        b=df1[df1.columns[-1]].idxmax()
        ml+=[b," segment is anticipated to reach ",str(round(df[df.columns[-2]][b],2)) ," ",value," by ", df.columns[-2]," with the highest CAGR of ",str(round(df[df.columns[-1]][b],2)),"%. "]
        df2=df1.drop(index=[a,b])
        c=df2[df2.columns[0]].idxmax()
        d=df2[df2.columns[-1]].idxmax()
        e=round(((df[df.columns[0]][a]+df[df.columns[0]][c])/df[df.columns[0]]["Total"])*100,1)
        ml+=[a," and ",c," segments collectively expected to account for about ",str(e),"% share of the ",name," market in ",df.columns[0],", with the former constituting around ",str(round((df[df.columns[0]][a]/df[df.columns[0]]["Total"])*100,1)),"% share. ",b," and ",d," segments are expected to witness significant growth rates at a CAGR of ",str(round(df[df.columns[-1]][b],2)),"% and ",str(round(df[df.columns[-1]][d],2)),"% respectively, during the forecast period. Presently, share of these two segments is estimated to be around ",str(round(((df[df.columns[0]][b]+df[df.columns[0]][d])/df[df.columns[0]]["Total"])*100,2)),"% in the overall ",name," market in ",df.columns[0],", and is anticipated to reach ",str(round(((df[df.columns[-2]][b]+df[df.columns[-2]][d])/df[df.columns[-2]]["Total"])*100,2)),"% by ",df.columns[-2],". "]
        return("".join(ml))


def tocspl(toc,start,end):
    for cou in toc:
        for k1 in range (start,end):
            cou["split"].pop(str(k1))
            cou["vsplit"].pop(str(k1))

        for l1 in cou["toc"]:
            for l2 in l1["toc_component"]:
                for k1 in range (start,end):
                        l2["split"].pop(str(k1))
                        l2["vsplit"].pop(str(k1))
                if l2["toc2_component"]!=[]:
                    for l3 in l2["toc2_component"]:
                        for k1 in range (start,end):
                            l3["split"].pop(str(k1))
                            l3["vsplit"].pop(str(k1))
                        if l3["toc3_component"]!=[]:
                            for l4 in l3["toc3_component"]:
                                for k1 in range (start,end):
                                    l4["split"].pop(str(k1))
                                    l4["vsplit"].pop(str(k1))

def createreport(pdlist,market,reportname,value,volume,tablesplit,g):
    dir=os.getcwd()
    doc = docx.Document("Template NM _New.docx")
    # reportname=reportname.split("(1)")[0] 
    paragraph_format1 = doc.styles["Table1_Header"].paragraph_format
    paragraph_format2 = doc.styles['Table1_Left'].paragraph_format
    paragraph_format3 = doc.styles['Table1_Right'].paragraph_format
    paragraph_format4 = doc.styles['Table_No_NM'].paragraph_format
    paragraph_format5 = doc.styles['FIG_no_NM'].paragraph_format
    paragraph_format6 = doc.styles['Main_Heading-NM'].paragraph_format
    paragraph_format7 = doc.styles['Head2_NM'].paragraph_format
    paragraph_format8 = doc.styles['Head3_NM'].paragraph_format
    paragraph_format9 = doc.styles['head4_NM'].paragraph_format
    paragraph_format10 = doc.styles['Head5_NM'].paragraph_format
    paragraph_format11 = doc.styles['Head6_NM'].paragraph_format
#     paragraph_format12= doc.styles['Table1_Right'].paragraph_format
#     paragraph_format13 = doc.styles['Table_No_NM'].paragraph_format
#     paragraph_format14 = doc.styles['FIG_no_NM'].paragraph_format
    paragraph_format1.keep_with_next = True
    paragraph_format2.keep_with_next = True
    paragraph_format3.keep_with_next = True
    paragraph_format4.keep_with_next = True
    paragraph_format5.keep_with_next = True
    paragraph_format6.keep_with_next = True
    paragraph_format7.keep_with_next = True
    paragraph_format8.keep_with_next = True
    paragraph_format9.keep_with_next = True
    paragraph_format10.keep_with_next = True
    paragraph_format11.keep_with_next = True
    
    dt = datetime.datetime.now()
    t = str(datetime.datetime.timestamp(dt))
    # for k in pdlist:
    #     f=pd.DataFrame.from_dict(pdlist[k]["data"],orient='index')
    #     fvol=pd.DataFrame.from_dict(pdlist[k]["vdata"],orient='index')
    #     mt=(pdlist[k]["val"],", ",str(f.columns[0]),"-",str(f.columns[-1]),", ",value)
    #     mtvol=(pdlist[k]["vol"],", ",str(f.columns[0]),"-",str(f.columns[-1]),", ",volume)        
    #     if pdlist[k]["heading"]==1:
    #         doc.add_paragraph(k,style=parastyle(pdlist[k]["heading"]))
    #         doc.add_paragraph("Overview",style=parastyle(2))
    #     elif pdlist[k]["heading"]==3 and pdlist[k]["before"]!="":
    #         doc.add_paragraph(pdlist[k]["before"],style=parastyle(2))
    #         doc.add_paragraph(k,style=parastyle(pdlist[k]["heading"]))
    #     elif pdlist[k]["heading"]==5 and pdlist[k]["before"]!="":
    #         doc.add_paragraph(pdlist[k]["before"],style=parastyle(4))
    #         doc.add_paragraph(k,style=parastyle(pdlist[k]["heading"]))
    #     else:
    #         doc.add_paragraph(k,style=parastyle(pdlist[k]["heading"]))


    #     f.loc["Total"] = f.sum()
    #     f['CAGR % '+"("+f.columns[1]+"-"+f.columns[-1]+")" ] = ((f.iloc[:, -1].div(f.iloc[:, 1]).pow(1./(len(f.columns) - 1)).sub(1))*100 )    
        
    #     fvol.loc["Total"] = fvol.sum()
    #     fvol['CAGR % '+"("+fvol.columns[1]+"-"+fvol.columns[-1]+")" ] = ((fvol.iloc[:, -1].div(fvol.iloc[:, 1]).pow(1./(len(fvol.columns) - 1)).sub(1))*100 )
    Tabno = NamedStyle(name="Tab_no")
    Tabno.font = Font(name=' Franklin Gothic Medium Cond',
                    size=12,
                    bold=True,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')

    heading1 = NamedStyle(name="heading1")
    heading1.font = Font(name='Franklin Gothic Demi Cond',
                    size=19,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')


    heading2= NamedStyle(name="heading2")
    heading2.font = Font(name='Franklin Gothic Medium Cond',
                    size=14,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')


    source=NamedStyle(name="source")
    source.font=Font(name='Franklin Gothic Book',
                    size=7.5,
                    bold=False,
                    italic=True,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='545454')

    main_heading = NamedStyle(name="main_heading")
    main_heading.font = Font(name='Franklin Gothic Medium Cond',
                    size=35,
                    bold=False,
                    italic=False,
                    vertAlign=None,
                    underline='none',
                    strike=False,
                    color='1F497D')

    number=NamedStyle(name="number",number_format="0.00")


    wb = Workbook(write_only=True)
    ws = wb.create_sheet()
    ws.sheet_view.showGridLines = False

    cellmain = WriteOnlyCell(ws)
    cellmain.style = main_heading
    cellmain.alignment = Alignment(horizontal='center', vertical='center')  

    cell = WriteOnlyCell(ws)
    cell.style = Tabno

    cellh1 = WriteOnlyCell(ws)
    cellh1.style = heading1
        

    cellh2 = WriteOnlyCell(ws)
    cellh2.style = heading2

        
    cell1 = WriteOnlyCell(ws)
    cell1.style = source

    cellnum= WriteOnlyCell(ws)
    cellnum.style = number



    def format_row(row, cell):
        for c in row:
            cell.value = c
            yield cell

    styletable = TableStyleInfo(name="TableStyleMedium2", showFirstColumn=True,
                       showLastColumn=True, showRowStripes=True, showColumnStripes=False)
    
    tabname=1

    rnum =1      
    count=1
    count1=1
    count2=1

    lstcolmain=chr(ord('c')+(g[1]-g[0]+2))

    mainr=[None,None,market+reportname+" Market"]
    mainrow=format_row(mainr, cellmain)
    ws.append(mainrow)
    ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcolmain+str(rnum))
    rnum+=1

    col=[str(i) for i in range(g[0],g[1]+1)]
    for k in pdlist:
        try:

            f=pd.DataFrame.from_dict(pdlist[k]["data"],orient='index')
            f = f.reindex(columns=col)
            fvol=pd.DataFrame.from_dict(pdlist[k]["vdata"],orient='index')
            fvol = fvol.reindex(columns=col)
            mt=(k.split("#")[0],", ",f.columns[0],"-",f.columns[-1],", ",value)
            mtvol=(k.split("#")[0],", ",f.columns[0],"-",f.columns[-1],", ",volume)
            if pdlist[k]["heading"]==1:
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))
                doc.add_paragraph("Overview",style=parastyle(2))
            elif pdlist[k]["heading"]==3 and pdlist[k]["before"]!="":
                doc.add_paragraph(pdlist[k]["before"],style=parastyle(2))
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))
            elif pdlist[k]["heading"]==5 and pdlist[k]["before"]!="":
                doc.add_paragraph(pdlist[k]["before"],style=parastyle(4))
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))
            else:
                doc.add_paragraph(k.split("#")[0],style=parastyle(pdlist[k]["heading"]))

            
            # p=doc.add_paragraph("".join(mt),style="FIG_no_NM")
        
            f.loc["Total"] = f.sum()
            f['CAGR % '+"("+f.columns[1]+"-"+f.columns[-1]+")" ] = ((f.iloc[:, -1].div(f.iloc[:, 1]).pow(1./(len(f.columns) - 1)).sub(1))*100 ) 

            fvol.loc["Total"] = fvol.sum()
            fvol['CAGR % '+"("+fvol.columns[1]+"-"+fvol.columns[-1]+")" ] = ((fvol.iloc[:, -1].div(fvol.iloc[:, 1]).pow(1./(len(fvol.columns) - 1)).sub(1))*100 )
            if volume=="Units":
                fvol = fvol.round().astype({i:'int' for i in col })

            lstcol=chr(ord('c')+f.shape[1])
            if pdlist[k]["heading"]==1:
                h1=[None,None,str(count1)+". "+(k.split("#")[0]).upper()]
                count1+=1
                count2=1
                head_row = format_row(h1, cellh1)
                ws.append(head_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1  
                        
                
            elif pdlist[k]["heading"]==3 and pdlist[k]["before"]!="":
                h2=[None,None,str(count2)+". "+(pdlist[k]["before"]).upper()]
                count2+=1
                head2_row = format_row(h2, cellh2)
                ws.append(head2_row)
                ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
                rnum +=1  

                
                
                
            r1=[None,None,"TABLE "+str(count)+". "+("".join(mt)).upper()]
            count+=1
            row1=[None,None] 
            rows = (dataframe_to_rows(f, index=True, header=True))
            # first_row = format_first_row(row1+list(next(rows)), cell)
            first_row = format_row(r1, cell)
            ws.append(first_row)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            rnum +=1  
            rowtable=rnum
            df=list(dataframe_to_rows(f, index=True, header=True))
            for i in range (len(df)):
                row1=[None,None]
                row= row1+df[i]
            #         print((row))
                if i==0:
                    row=list(map(str, df[i]))
                    row[0]=pdlist[k]["comp"]
                    row=row1+row
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  
                elif i==1:
                    continue
                else:
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  

            tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
            tabname+=1
            headings = [pdlist[k]["comp"]]+list(f.columns)
            tab1._initialise_columns()
            for column, value1 in zip(tab1.tableColumns, headings):
                column.name = value1
        #     tab1.headerRowCount = False
            
            tab1.tableStyleInfo = styletable
            ws.add_table(tab1)
            sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
            sourcerow = format_row(sourcer, cell1)
            ws.append(sourcerow)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            ws.append([None])
            ws.append([None])
            rnum +=3

            r1=[None,None,"TABLE "+str(count)+". "+("".join(mtvol)).upper()]
            count+=1
            row1=[None,None] 
            rows = (dataframe_to_rows(fvol, index=True, header=True))
            # first_row = format_first_row(row1+list(next(rows)), cell)
            first_row = format_row(r1, cell)
            ws.append(first_row)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            rnum +=1 
            rowtable=rnum 
            df=list(dataframe_to_rows(f, index=True, header=True))
            for i in range (len(df)):
                row1=[None,None]
                row= row1+df[i]
            #         print((row))
                if i==0:
                    row=list(map(str, df[i]))
                    row[0]=pdlist[k]["comp"]
                    row=row1+row
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  
                elif i==1:
                    continue
                else:
                    roe_num=format_row(row,cellnum)
                    ws.append(roe_num)
                    rnum +=1  

            tab1 = Table(displayName="table"+str(tabname), ref="c"+str(rowtable)+":"+lstcol+str(rnum-1))
            tabname+=1
            headings = [pdlist[k]["comp"]]+list(f.columns)
            tab1._initialise_columns()
            for column, value1 in zip(tab1.tableColumns, headings):
                column.name = value1
        #     tab1.headerRowCount = False
            
            tab1.tableStyleInfo = styletable
            ws.add_table(tab1)
            sourcer=[None,None,"Source: Primary Research, Secondary Research, NMSC Analysis"]
            sourcerow = format_row(sourcer, cell1)
            ws.append(sourcerow)
            ws.merged_cells.ranges.append("c"+str(rnum)+":"+lstcol+str(rnum))
            ws.append([None])
            ws.append([None])
            rnum +=3





            if (tablesplit):
                f=f.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)
                fvol=fvol.drop([str(i) for i in range (g[0]+3,g[1],2)],axis=1)

            p=doc.add_paragraph("".join(mt),style="FIG_no_NM")
            f1=f.drop([f.columns[-1]],axis=1)
            f1=f1.drop(index=("Total"))
            pic=plot(f1.T,t)
            doc.add_picture(pic,width=Cm(16.98), height=Cm(8.51))
            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")
            # doc.add_paragraph(" ")



            p=doc.add_paragraph("".join(mt),style="Table_No_NM")
            table = doc.add_table(rows=f.shape[0]+1, cols=f.shape[1]+1)
            table.style = 'Table_Style_NM'
            table_cells = table._cells
            for r in range (len(table.rows)):
                row=table.rows[r]
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                if r==0:
                    row.height = Cm(1.1)
                else:
                    row.height = Cm(0.7)

            table.cell(0,0).paragraphs[0].style="Table1_Header"
            table.cell(0,0).width = 1242000.2592
            table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
            table.columns[0].width= Cm(3.45)
            table.columns[-1].width=Cm(1.97)


            for k1 in range(1,len(f.columns)+1):
                table.cell(0,k1).paragraphs[0].style="Table1_Header"
                table.cell(0,k1).paragraphs[0].text=(f.columns[k1-1])
            for i in range(1,f.shape[0]+1):
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].style="Table1_Left"
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].text=(f.index[i-1])
                for j in range(1,f.shape[1]+1):
                    table_cells[j + i * (f.shape[1]+1)].paragraphs[0].style="Table1_Right"
                    if (j %(f.shape[1]))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2))+"%")
                        
                    elif (j %(f.shape[1]+1))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 1242000.2592
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))
                        
                    else:
                        table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))

            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis",style="Source_NM")
            

            doc.add_paragraph(writeup(f,pdlist[k]["wrtup"],value),style="Normal")

            
            p=doc.add_paragraph("".join(mtvol),style="FIG_no_NM")
            f1vol=fvol.drop([fvol.columns[-1]],axis=1)
            f1vol=f1vol.drop(index=("Total"))
            picv=plot(f1vol.T,t)
            doc.add_picture(picv,width=Cm(16.98), height=Cm(8.51))
            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis \n",style="Source_NM")
            # doc.add_paragraph(" ")

            
            p=doc.add_paragraph("".join(mtvol),style="Table_No_NM")
            table = doc.add_table(rows=fvol.shape[0]+1, cols=fvol.shape[1]+1)
            table.style = 'Table_Style_NM'
            table_cells = table._cells
            for r in range (len(table.rows)):
                row=table.rows[r]
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
                if r==0:
                    row.height = Cm(1.1)
                else:
                    row.height = Cm(0.7)

            table.cell(0,0).paragraphs[0].style="Table1_Header"
            table.cell(0,0).width = 1242000.2592
            table.cell(0,0).paragraphs[0].text=pdlist[k]["comp"]
            table.columns[0].width= Cm(3.45)
            table.columns[-1].width=Cm(1.97)

            for k1 in range(1,len(f.columns)+1):
                table.cell(0,k1).paragraphs[0].style="Table1_Header"
                table.cell(0,k1).paragraphs[0].text=(f.columns[k1-1])
            for i in range(1,f.shape[0]+1):
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].style="Table1_Left"
                table_cells[ i*(f.shape[1]+1)].paragraphs[0].text=(f.index[i-1])
                for j in range(1,f.shape[1]+1):
                    table_cells[j + i * (f.shape[1]+1)].paragraphs[0].style="Table1_Right"
                    if (j %(f.shape[1]))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 709200.04464
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2))+"%")
                        
                    elif (j %(f.shape[1]+1))==0:
                        table_cells[j + i * (f.shape[1]+1)].width = 1242000.2592
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))
                        
                    else:
                        table_cells[j + i * (f.shape[1]+1)].width = 594000.00072
                        table_cells[j + i * (f.shape[1]+1)].paragraphs[0].text=(str(round(f.values[i-1][j-1],2)))

            doc.add_paragraph("Source: Primary Research, Secondary Research, NMSC Analysis",style="Source_NM")
           
            doc.add_paragraph(writeupvol(fvol,pdlist[k]["wrtup"],volume),style="Normal")

        except Exception as e:
                dir= os.getcwd()
                logfile=dir+"\error_log.txt"
                file=open(logfile,"a")
                file.write("/n"+market.upper()+" "+reportname.upper()+" "+str(e))
                print(e)

                continue

    doc.save(dir+"/files/"+market+" "+reportname+" Value and Volume Market Report_Forecast Till "+str(g[1])+".docx")
    wb.save(dir+"/files/"+market+" "+reportname+" Value and Volume Market Report_Forecast Excel Till "+str(g[1])+".xlsx")
    os.remove(pic)
    # os.remove(picv)
    return (("/files/"+market+" "+reportname+" Value and Volume Market Report_Forecast Till "+str(g[1])+".docx","/files/"+market+" "+reportname+" Value and Volume Market Report_Forecast Excel Till "+str(g[1])+".xlsx"))


def reportmake(toc1,year,ASP,market,reportname,value,volume,tablesplit):
    v=list(year.keys())
    # yearvol={i:mpf(year[i])/mpf(ASP[i]) for i in year}
    reportname=reportname.split("(1)")[0]
    for i in year:
        sk=year[i]
        year[i]=sk.replace(" ","")
        year[i]=sk.replace(",","")
    for n in ASP:
        sk=ASP[n]
        ASP[n]=sk.replace(" ","")
        ASP[n]=sk.replace(",","")

    if volume=="Units":
        yearvol={i:int(round(float(ASP[i]))) for i in ASP}
    else:
        yearvol=ASP
    g=[int(v[0]),int(v[-1])]
    region100(toc1)
    toc100(toc1)
    toc1=tocchanger(toc1,year,yearvol,g)
    # print(toc1)
    f=det(toc1)
    pdlist=pdlistmaker(toc1,f[1],f[0],market,reportname)
    contchange=[{"label":j,"country":j,"value":j} for j in f[0]]
    return(createreport(pdlist,market,reportname,value,volume,tablesplit,g),{"region":[market],"country":f[0]},{market:contchange})

def reportmakedownedit(toc1,region,market,reportname,value,volume,startyear,endyear,tablesplit):
    print(region)
    reportname=reportname.split("(1)")[0]
    for i in toc1:
        if i['region']==region:
            toc1=i["tocreg"]
    f=det(toc1)
    g=[startyear,endyear]
    pdlist=pdlistmaker(toc1,f[1],f[0],market,reportname)
    return(createreport(pdlist,market,reportname,value,volume,tablesplit,g))


def reportmakedownedityear(toc1,region,market,reportname,value,volume,tablesplit,year,startyear,endyear,oggeo):
    reportname=reportname.split("(1)")[0]
    if oggeo=="Global":
        for i in toc1:
            if i['region']==region:
                toc1=i["tocreg"]
    
    f=det(toc1)
    tocspl(toc1,startyear,year)
    g=[year,endyear]
    pdlist=pdlistmaker(toc1,f[1],f[0],market,reportname)
    return(createreport(pdlist,market,reportname,value,volume,tablesplit,g))